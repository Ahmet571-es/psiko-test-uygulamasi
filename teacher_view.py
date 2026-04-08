import streamlit as st
# PERFORMANS: pandas, matplotlib, seaborn, openpyxl, pdf_engine
# sadece kullanıldıkları fonksiyonların içinde import edilir.
# Bu sayede öğretmen paneli açılırken 5-10 sn tasarruf sağlanır.
import json
import time
import os
import io
from datetime import datetime
from dotenv import load_dotenv
from db_utils import (
    get_all_students_with_results, reset_database,
    delete_specific_students, save_holistic_analysis,
    get_student_analysis_history, is_using_sqlite
)

# --- API AYARLARI ---
load_dotenv()

# Claude model — ortam değişkeni veya Streamlit secrets ile değiştirilebilir
DEFAULT_MODEL = "claude-sonnet-4-20250514"
def _get_claude_model():
    try:
        if "CLAUDE_MODEL" in st.secrets:
            return st.secrets["CLAUDE_MODEL"]
    except Exception:
        pass
    return os.getenv("CLAUDE_MODEL", DEFAULT_MODEL)


def get_claude_client():
    """
    Claude API istemcisini başlatır.
    Öncelik: st.secrets → .env dosyası → ortam değişkeni
    """
    try:
        if "ANTHROPIC_API_KEY" in st.secrets:
            api_key = st.secrets["ANTHROPIC_API_KEY"]
        else:
            api_key = os.getenv("ANTHROPIC_API_KEY")

        if not api_key:
            return None

        from anthropic import Anthropic
        return Anthropic(api_key=api_key)
    except ImportError:
        return None
    except Exception:
        return None


# --- AI ANALİZ FONKSİYONU ---
def get_ai_analysis(prompt):
    """Claude API ile analiz üretir."""
    client = get_claude_client()
    if not client:
        return "⚠️ Hata: Claude API Key bulunamadı veya 'anthropic' kütüphanesi eksik. Lütfen Streamlit Secrets veya .env dosyasını kontrol edin."

    try:
        response = client.messages.create(
            model=_get_claude_model(),
            max_tokens=16000,
            temperature=0.3,
            messages=[{"role": "user", "content": prompt}]
        )
        return response.content[0].text

    except Exception as e:
        err = str(e)
        if "authentication" in err.lower():
            return "⚠️ API Key hatalı veya geçersiz. Lütfen Streamlit Secrets'taki ANTHROPIC_API_KEY değerini kontrol edin."
        elif "invalid_request" in err.lower() or "model" in err.lower():
            return f"⚠️ Model hatası: {err}"
        elif "rate_limit" in err.lower():
            return "⚠️ API istek limiti aşıldı. Lütfen birkaç dakika bekleyip tekrar deneyin."
        else:
            return f"⚠️ Analiz sırasında bir hata oluştu: {err}"


# --- GRAFİK FONKSİYONU (aşağıda tanımlı) ---

# ============================================================
# EXCEL VERİ DIŞA AKTARMA FONKSİYONLARI
# ============================================================

def _style_header(ws, row, max_col):
    """Başlık satırını biçimlendir."""
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    header_font = Font(name='Arial', bold=True, color='FFFFFF', size=11)
    header_fill = PatternFill('solid', fgColor='1B2A4A')
    header_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    for col in range(1, max_col + 1):
        cell = ws.cell(row=row, column=col)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = thin_border


def _style_data_rows(ws, start_row, end_row, max_col):
    """Veri satırlarını biçimlendir."""
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    data_font = Font(name='Arial', size=10)
    data_align = Alignment(vertical='top', wrap_text=True)
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    zebra_fill = PatternFill('solid', fgColor='F2F4F8')
    for row in range(start_row, end_row + 1):
        for col in range(1, max_col + 1):
            cell = ws.cell(row=row, column=col)
            cell.font = data_font
            cell.alignment = data_align
            cell.border = thin_border
            if row % 2 == 0:
                cell.fill = zebra_fill


def _auto_width(ws, max_col, max_width=60):
    """Sütun genişliklerini otomatik ayarla."""
    from openpyxl.utils import get_column_letter
    for col in range(1, max_col + 1):
        max_len = 0
        for row in ws.iter_rows(min_col=col, max_col=col, values_only=False):
            for cell in row:
                if cell.value:
                    lines = str(cell.value).split('\n')
                    max_len = max(max_len, max(len(l) for l in lines))
        adjusted = min(max(max_len + 2, 12), max_width)
        ws.column_dimensions[get_column_letter(col)].width = adjusted


def generate_full_system_excel(all_data):
    """Tüm sistem verisini çok sayfalı Excel dosyasına dönüştürür."""
    from openpyxl import Workbook
    wb = Workbook()

    # ── SAYFA 1: ÖĞRENCİ LİSTESİ ──
    ws1 = wb.active
    ws1.title = "Öğrenci Listesi"
    headers1 = ["No", "Ad Soyad", "E-posta", "Yaş", "Cinsiyet", "Sınıf", "Giriş Sayısı", "Çözülen Test"]
    ws1.append(headers1)
    _style_header(ws1, 1, len(headers1))

    for i, d in enumerate(all_data, 1):
        info = d["info"]
        grade_text = f"{info.grade}. Sınıf" if info.grade else "—"
        ws1.append([
            i, info.name, info.username, info.age, info.gender,
            grade_text, info.login_count, len(d["tests"])
        ])
    _style_data_rows(ws1, 2, len(all_data) + 1, len(headers1))
    _auto_width(ws1, len(headers1))
    ws1.freeze_panes = 'A2'

    # ── SAYFA 2: TEST SONUÇLARI ──
    ws2 = wb.create_sheet("Test Sonuçları")
    headers2 = ["No", "Öğrenci Adı", "Test Adı", "Tarih", "Puan Detayları", "Sistem Raporu"]
    ws2.append(headers2)
    _style_header(ws2, 1, len(headers2))

    row_idx = 0
    for d in all_data:
        for t in d["tests"]:
            row_idx += 1
            scores_text = ""
            if t["scores"]:
                if isinstance(t["scores"], dict):
                    scores_text = "\n".join(f"{k}: {v}" for k, v in t["scores"].items())
                else:
                    scores_text = str(t["scores"])
            report_text = t.get("report", "") or ""
            if len(report_text) > 32000:
                report_text = report_text[:32000] + "\n... (kesildi)"
            ws2.append([
                row_idx, d["info"].name, t["test_name"],
                str(t["date"]), scores_text, report_text
            ])
    _style_data_rows(ws2, 2, row_idx + 1, len(headers2))
    _auto_width(ws2, len(headers2))
    ws2.freeze_panes = 'A2'

    # ── SAYFA 3: AI ANALİZ RAPORLARI ──
    ws3 = wb.create_sheet("AI Analiz Raporları")
    headers3 = ["No", "Öğrenci Adı", "Test Kombinasyonu", "Rapor Tarihi", "AI Rapor İçeriği"]
    ws3.append(headers3)
    _style_header(ws3, 1, len(headers3))

    ai_row = 0
    for d in all_data:
        history = get_student_analysis_history(d["info"].id)
        for h in history:
            ai_row += 1
            report_content = h.get("report", "") or ""
            if len(report_content) > 32000:
                report_content = report_content[:32000] + "\n... (kesildi)"
            ws3.append([
                ai_row, d["info"].name, h["combination"],
                str(h["date"]), report_content
            ])
    if ai_row == 0:
        ws3.append(["", "Henüz AI raporu oluşturulmamış.", "", "", ""])
        ai_row = 1
    _style_data_rows(ws3, 2, ai_row + 1, len(headers3))
    _auto_width(ws3, len(headers3))
    ws3.freeze_panes = 'A2'

    # ── SAYFA 4: DETAYLI PUAN MATRİSİ ──
    ws4 = wb.create_sheet("Puan Matrisi")
    all_score_keys = set()
    for d in all_data:
        for t in d["tests"]:
            if t["scores"]:
                all_score_keys.update(t["scores"].keys())
    all_score_keys = sorted(all_score_keys)

    headers4 = ["Öğrenci Adı", "Test Adı", "Tarih"] + all_score_keys
    ws4.append(headers4)
    _style_header(ws4, 1, len(headers4))

    matrix_row = 0
    for d in all_data:
        for t in d["tests"]:
            if t["scores"]:
                matrix_row += 1
                row_data = [d["info"].name, t["test_name"], str(t["date"])]
                for key in all_score_keys:
                    val = t["scores"].get(key, "")
                    if isinstance(val, (dict, list)):
                        val = str(val)
                    elif isinstance(val, bool):
                        val = str(val)
                    elif val is None:
                        val = ""
                    elif not isinstance(val, (int, float, str)):
                        val = str(val)
                    row_data.append(val)
                ws4.append(row_data)
    if matrix_row == 0:
        ws4.append(["Henüz puan verisi yok"] + [""] * (len(headers4) - 1))
        matrix_row = 1
    _style_data_rows(ws4, 2, matrix_row + 1, len(headers4))
    _auto_width(ws4, len(headers4))
    ws4.freeze_panes = 'A2'

    # Buffer'a yaz
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer


def generate_student_excel(student_data, analysis_history):
    """Tek öğrenci verisini Excel dosyasına dönüştürür."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    info = student_data["info"]
    tests = student_data["tests"]
    wb = Workbook()

    # ── SAYFA 1: ÖĞRENCİ PROFİLİ ──
    ws1 = wb.active
    ws1.title = "Öğrenci Profili"
    title_font = Font(name='Arial', bold=True, size=14, color='1B2A4A')
    label_font = Font(name='Arial', bold=True, size=11)
    value_font = Font(name='Arial', size=11)

    ws1.merge_cells('A1:B1')
    ws1['A1'] = f"📁 {info.name} — Öğrenci Dosyası"
    ws1['A1'].font = title_font

    profile_data = [
        ("Ad Soyad", info.name),
        ("E-posta", info.username),
        ("Yaş", info.age),
        ("Cinsiyet", info.gender),
        ("Sınıf", f"{info.grade}. Sınıf" if info.grade else "Belirtilmemiş"),
        ("Toplam Giriş", info.login_count),
        ("Çözülen Test", len(tests)),
        ("Rapor Tarihi", datetime.now().strftime("%d.%m.%Y %H:%M")),
    ]
    for i, (label, value) in enumerate(profile_data, 3):
        ws1.cell(row=i, column=1, value=label).font = label_font
        ws1.cell(row=i, column=2, value=value).font = value_font

    ws1.column_dimensions['A'].width = 20
    ws1.column_dimensions['B'].width = 40

    # ── SAYFA 2: TEST SONUÇLARI ──
    ws2 = wb.create_sheet("Test Sonuçları")
    headers2 = ["No", "Test Adı", "Tarih", "Puan Detayları", "Sistem Raporu"]
    ws2.append(headers2)
    _style_header(ws2, 1, len(headers2))

    for i, t in enumerate(tests, 1):
        scores_text = ""
        if t["scores"]:
            scores_text = "\n".join(f"{k}: {v}" for k, v in t["scores"].items())
        report_text = t.get("report", "") or ""
        if len(report_text) > 32000:
            report_text = report_text[:32000] + "\n... (kesildi)"
        ws2.append([i, t["test_name"], t["date"], scores_text, report_text])

    _style_data_rows(ws2, 2, len(tests) + 1, len(headers2))
    _auto_width(ws2, len(headers2))
    ws2.freeze_panes = 'A2'

    # ── SAYFA 3: AI RAPORLARI ──
    ws3 = wb.create_sheet("AI Raporları")
    headers3 = ["No", "Test Kombinasyonu", "Rapor Tarihi", "AI Rapor İçeriği"]
    ws3.append(headers3)
    _style_header(ws3, 1, len(headers3))

    for i, h in enumerate(analysis_history, 1):
        report_content = h.get("report", "") or ""
        if len(report_content) > 32000:
            report_content = report_content[:32000] + "\n... (kesildi)"
        ws3.append([i, h["combination"], h["date"], report_content])

    ai_count = len(analysis_history) if analysis_history else 0
    if ai_count == 0:
        ws3.append(["", "Henüz AI raporu oluşturulmamış.", "", ""])
        ai_count = 1
    _style_data_rows(ws3, 2, ai_count + 1, len(headers3))
    _auto_width(ws3, len(headers3))
    ws3.freeze_panes = 'A2'

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer


# ============================================================
# GRAFİK FONKSİYONU
# ============================================================

def plot_scores(data_dict, title):
    """Test sonuçlarını görselleştirmek için Bar Grafiği oluşturur."""
    import matplotlib.pyplot as plt
    import seaborn as sns

    if not data_dict or not isinstance(data_dict, dict):
        return None

    plot_data = {}

    # 1. Durum: 'categories' anahtarı varsa (Çalışma Davranışı, Sınav Kaygısı)
    if "categories" in data_dict and isinstance(data_dict["categories"], dict):
        plot_data = data_dict["categories"]

    # 1b. Durum: P2 Dikkat Testi
    elif "CP" in data_dict and "TN_E" in data_dict:
        d2_labels = {
            "CP": "Konsantrasyon (CP)",
            "TN_E": "Toplam Performans (TN-E)",
            "TN": "Toplam İşaretleme (TN)",
            "E1": "Atlama Hatası (E1)",
            "E2": "Yanlış İşaretleme (E2)",
            "FR": "Dalgalanma (FR)",
        }
        for key, label in d2_labels.items():
            if key in data_dict and isinstance(data_dict[key], (int, float)):
                plot_data[label] = data_dict[key]

    # 1c. Durum: Akademik Analiz Testi
    elif "overall" in data_dict and "performance_avg" in data_dict:
        akd_keys = {
            "overall": "Genel Skor",
            "Anlama": "Okuma Anlama",
            "Muhakeme": "Matematiksel Muhakeme",
            "Düşünme": "Mantıksal Düşünme",
            "Öz-Değerlendirme": "Öz-Değerlendirme",
        }
        for key, label in akd_keys.items():
            if key in data_dict and isinstance(data_dict[key], (int, float)):
                plot_data[label] = data_dict[key]

    # 2. Durum: 'scores' anahtarı varsa (Çoklu Zeka)
    elif "scores" in data_dict and isinstance(data_dict["scores"], dict):
        temp_data = {}
        for k, v in data_dict["scores"].items():
            if isinstance(v, dict) and "pct" in v:
                temp_data[k] = v["pct"]
            elif isinstance(v, (int, float)):
                temp_data[k] = v
        plot_data = temp_data if temp_data else data_dict["scores"]

    # 3. Durum: Düz sözlük (Sağ-Sol Beyin, VARK, Holland)
    else:
        for k, v in data_dict.items():
            if not isinstance(v, (int, float)):
                continue
            if k in ["id", "user_id", "total", "max_total", "total_responses", "total_pct"]:
                continue
            if "yuzde" in k:
                label = "Sağ Beyin %" if "sag" in k else "Sol Beyin %"
                plot_data[label] = v
            elif k in ["beyin", "dominant", "level", "version"]:
                continue
            elif k in ["sag_beyin", "sol_beyin"]:
                continue
            else:
                plot_data[k] = v

    if not plot_data:
        return None

    # Veriyi hazırla
    valid_pairs = [(str(k), float(v)) for k, v in plot_data.items() if isinstance(v, (int, float))]
    if not valid_pairs:
        return None

    labels = [p[0] for p in valid_pairs]
    values = [p[1] for p in valid_pairs]

    # Grafik Ayarları
    sns.set_theme(style="whitegrid")
    fig, ax = plt.subplots(figsize=(8, max(3, len(labels) * 0.5)))

    # Renk paleti — logo uyumlu
    colors = sns.color_palette("coolwarm", len(labels))
    sns.barplot(x=values, y=labels, hue=labels, ax=ax, palette=colors, orient='h', legend=False)

    ax.set_title(f"{title}", fontsize=12, fontweight='bold', color='#1B2A4A')
    ax.set_xlabel("Puan / Yüzde", fontsize=10)
    ax.tick_params(axis='y', labelsize=9)
    plt.tight_layout()
    return fig


# ============================================================
# PROMPT ÜRETME FONKSİYONLARI — TİCARİ KALİTE v3.0
# ============================================================

def build_holistic_prompt(student_name, student_age, student_gender, test_data_list, student_grade=None):
    """Bütüncül (harmanlanmış) analiz için ticari kalite prompt."""
    grade_text = f"{student_grade}. Sınıf" if student_grade else "Belirtilmemiş"
    return f"""# ROL ve KİMLİK

Sen, Türkiye'nin önde gelen eğitim psikolojisi merkezlerinde 20 yıl deneyim kazanmış, psikometrik değerlendirme, kariyer danışmanlığı ve gelişim psikolojisi alanlarında uzmanlaşmış bir Klinik Eğitim Psikoloğusun. 

Uzmanlık alanların:
- Psikometrik test bataryası yorumlama ve çapraz korelasyon analizi
- Ergen gelişim psikolojisi ve yaşa özgü değerlendirme
- Kariyer psikolojisi ve mesleki yönlendirme
- Aile danışmanlığı ve ebeveyn rehberliği
- Öğrenme farklılıkları ve bireyselleştirilmiş eğitim planlaması

Bu rapor, ücretli bir profesyonel danışmanlık hizmetinin çıktısıdır. Aile, öğretmen ve rehber öğretmenler tarafından okunacak resmi bir analiz belgesidir. Raporun, yüz yüze bir psikolog görüşmesinin yazılı karşılığı kadar derinlikli, kişiselleştirilmiş ve uygulanabilir olmalıdır.

---

# ÖĞRENCİ DOSYASI

| Alan | Bilgi |
|------|-------|
| İsim | {student_name} |
| Yaş | {student_age} |
| Sınıf | {grade_text} |
| Cinsiyet | {student_gender} |
| Değerlendirme Türü | Bütüncül Çoklu Test Analizi |

## TEST VERİLERİ (JSON)
```json
{json.dumps(test_data_list, ensure_ascii=False, indent=2)}
```

---

# KRİTİK KURALLAR

1. **KANITSAL ZORUNLULUK:** Her yorum, iddia ve tespit mutlaka parantez içinde kaynak test adı ve sayısal puan ile desteklenmeli. Örn: "Görsel-uzamsal zeka alanında belirgin güç göstermektedir (Çoklu Zeka: %82)." Kanıtsız hiçbir yorum yapma.

2. **SENTEZ MERKEZLİ:** Testleri ayrı ayrı özetleme. Asıl değer, testler arasındaki BAĞLANTILARDA, KORELASYONLARDA ve ÇELİŞKİLERDE yatıyor. Her paragrafta en az 2 farklı testten veri çaprazla.

3. **GELİŞİMSEL BAĞLAM:** {student_age} yaşında, {grade_text} düzeyinde bir bireyin gelişimsel dönem özelliklerini (bilişsel, duygusal, sosyal, kimlik gelişimi) göz önünde bulundurarak yorumla. Yaşa ve sınıf seviyesine özgü beklentileri ve normları referans al.

4. **TIBBİ TANI YASAĞI:** "DEHB", "depresyon", "anksiyete bozukluğu", "otizm spektrumu", "disleksi" gibi klinik tanı terimleri kesinlikle kullanma. Bunun yerine davranışsal betimleme yap.

5. **BİREYSELLEŞTİRME:** Genel geçer tavsiyeler verme. Her öneri, bu öğrencinin spesifik veri profilinden türetilmiş olmalı. "Daha çok çalış" yerine "VARK Kinestetik baskınlığın (%X) göz önüne alındığında, Matematik çalışırken manipülatif materyaller (geometri blokları, kağıt katlama) kullanman, masa başı süresini 25 dakikalık bloklara bölmen önerilir."

6. **PUAN YORUMLAMA ÇERÇEVESİ:**
   - %0-20 → Belirgin gelişim alanı — acil destek önerilir
   - %21-40 → Ortalamanın altı — hedefli çalışma gerektirir
   - %41-60 → Ortalama düzey — potansiyel mevcut, strateji ile yükseltilebilir
   - %61-80 → Güçlü alan — sürdürülebilir ve derinleştirilebilir
   - %81-100 → Çok güçlü / baskın alan — yetenek göstergesi, ileri düzey destekle parlayabilir

7. **UZUNLUK ve DERİNLİK:** Bu rapor minimum 3000 kelime olmalıdır. Her bölüm, ödenen ücrete değecek derinlikte olmalı. Yüzeysel veya şablonik ifadelerden kaçın. Her öğrenci için rapor benzersiz ve kişiselleştirilmiş olmalı.

8. **BİLİMSEL TEMEL ZORUNLULUĞU:** Tüm yorumlar ve tespitler, eğitim psikolojisi ve psikometri alanlarındaki bilimsel araştırmalara dayalı olmalıdır. Her önemli tespit için mümkünse teorik arka planı (ör. Gardner'ın Çoklu Zekâ Kuramı, Holland'ın Mesleki Kişilik Tipleri, Yerkes-Dodson Yasası) kısaca belirt. Ancak her bilimsel terim veya kavram kullanıldığında, hemen ardından parantez içinde lise öğrencisinin bile anlayabileceği yalın bir Türkçe açıklama ekle. Örnek: "Yerkes-Dodson Yasası'na göre (yani belirli bir noktaya kadar kaygı performansı artırır ama bu nokta aşılırsa performans düşmeye başlar)..." veya "Metakognitif farkındalık (kendi düşünme sürecini izleyebilme ve yönetebilme becerisi) düzeyi..."

9. **DENGELİ ve GERÇEKÇİ MOTİVASYON TONU:** Abartılı motivasyon ifadelerinden, aşırı iyimser veya pohpohlayıcı bir dilden kesinlikle kaçın. "Muhteşem", "olağanüstü", "inanılmaz potansiyel", "her şeyi başarabilir" gibi abartılı ifadeler YASAK. Bunun yerine gerçekçi, dengeli ve yapıcı bir ton kullan. Güçlü yönleri belirtirken somut veriye dayan, gelişim alanlarını görmezden gelme. Normal ve samimi bir motivasyon dili kullan: "Bu alanda güçlü bir profil ortaya koyuyor" gibi veriye dayalı, ölçülü ifadeler tercih et. Rapor, bir danışmanın sakin ve profesyonel üslubuyla yazılmalı — ne karamsar ne de abartılı iyimser.

10. **YALIN TÜRKÇE ZORUNLULUĞU:** Raporun tamamı, okuryazar bir lise öğrencisinin rahatlıkla anlayabileceği düzeyde yazılmalıdır. Akademik jargon ve teknik terimler kullanıldığında mutlaka parantez içinde kısa, günlük dile yakın bir açıklama ekle. Uzun ve karmaşık cümlelerden kaçın, kısa ve net cümleler kur. Yabancı kökenli sözcükler yerine mümkünse Türkçe karşılıklarını tercih et.

---

# TESTE ÖZEL ÇAPRAZ ANALİZ REHBERİ

Verideki test kombinasyonlarına göre aşağıdaki çapraz analizleri MUTLAKA yap:

## Enneagram Verisi Varsa:
- Ana tipin motivasyon yapısını diğer tüm test sonuçlarıyla çapraz kontrol et
- Kanat (wing) etkisinin öğrenme stili üzerindeki yansımasını VARK/Beyin dominansı ile doğrula
- Stres yönündeki tipin puanını Sınav Kaygısı verileriyle karşılaştır (stres tipi yüksekse kaygı da yüksek mi?)
- Büyüme yönündeki tipin puanını akademik güçlü alanlarla eşleştir
- Tritype analizi (Kafa 5-6-7 / Kalp 2-3-4 / Karın 8-9-1 merkezlerinden en yüksek puan) yap ve bütünsel kişilik portresini çiz
- Kişilik tipi ile Holland RIASEC kodu arasındaki uyumu/uyumsuzluğu tartış (Örn: Tip 5 + Araştırmacı(I) = uyumlu; Tip 5 + Girişimci(E) = çelişki)

## Sınav Kaygısı + Çalışma Davranışı Birlikte Varsa:
- Kaygı-performans döngüsünü analiz et: yetersiz çalışma → kaygı mı, yoksa kaygı → çalışamama mı?
- Hangi kaygı alt boyutu hangi çalışma davranışı kategorisiyle ilişkili?
- Bu döngüyü kırmak için somut müdahale noktasını tespit et

## VARK + Sağ-Sol Beyin Birlikte Varsa:
- "Nörobilişsel Öğrenme Profili" oluştur: beyin yarım küre baskınlığı + duyusal kanal tercihi
- Bu kombinasyonun sınıf ortamındaki optimal öğrenme koşullarını tanımla
- Ders bazlı (Matematik, Fen, Türkçe, Sosyal, Yabancı Dil) öğrenme stratejileri tablosu oluştur

## Çoklu Zeka + Holland RIASEC Birlikte Varsa:
- Zeka profili ile mesleki ilgi alanlarının örtüşme haritasını çıkar
- Uyumlu alanlar: doğal kariyer yönelimleri
- Uyumsuz alanlar: keşfedilmemiş potansiyel mi, yoksa yüzeysel ilgi mi?
- Top 10 kariyer önerisi (zeka + ilgi + kişilik üçgeninden)

---

# RAPOR FORMATI (HER BÖLÜMÜ AYNEN DOLDUR, HİÇBİR BÖLÜMÜ ATLAMA)

---

# 📋 YÖNETİCİ ÖZETİ

*(Bu bölüm, raporu okuyacak kişinin ilk 2 dakikada tüm tabloyu görmesini sağlar. 5-6 cümle ile öğrencinin en kritik güçlü yönü, en acil gelişim alanı, en dikkat çekici çelişki ve en öncelikli adım özetlenir.)*

---

# 🧬 1. KİŞİLİK ve MOTİVASYON PROFİLİ

## 1.1 Kim Bu Öğrenci?
*(Öğrenciyi hiç tanımayan birinin okuduğunda zihninde net bir portre oluşturacağı, 2-3 paragraflık derinlikli giriş. Tüm test verilerinden sentezlenmiş bir "karakter taslağı". Öğrencinin tipik bir gününü, sınıf davranışını, arkadaş ilişkilerini ve motivasyon kaynaklarını betimle.)*

## 1.2 Temel Motivasyon Dinamikleri
*(Bu öğrenci neyin peşinde koşuyor? Neyden kaçınıyor? Ne zaman en verimli? Ne zaman engellenmiş hissediyor? Enneagram + diğer test verileriyle desteklenmiş derinlikli motivasyon analizi. Minimum 2 paragraf.)*

## 1.3 Stres Tepki Profili
*(Bu öğrenci baskı altında nasıl tepki verir? Hangi durumlar tetikleyici? Kaçınma mı gösterir, aşırı çalışma mı, içe kapanma mı? Enneagram stres yönü + Sınav Kaygısı verileriyle destekle. Minimum 2 paragraf.)*

## 1.4 Sosyal ve Duygusal Harita
*(Akran ilişkileri, grup içi rolü, otorite figürleriyle ilişkisi, empati kapasitesi, çatışma yönetimi tarzı. Kişilik profili + sosyal/kişilerarası zeka verilerinden çıkarım. Minimum 2 paragraf.)*

---

# 🧠 2. BİLİŞSEL ve AKADEMİK PROFİL

## 2.1 Nörobilişsel Öğrenme Kimliği
*(Sağ/Sol Beyin dominansı + VARK öğrenme stili sentezi. Bu öğrencinin beyninin bilgiyi nasıl aldığını, işlediğini ve depoladığını açıkla. "Bu öğrenci bilgiyi önce GÖRÜR, sonra İŞLER, sonra HAREKET ile pekiştirir" gibi somut bir öğrenme akışı tanımla. Minimum 2 paragraf.)*

## 2.2 Zeka Profili Haritası
*(Çoklu Zeka verilerini detaylı yorumla. Profil tipi: uzmanlaşmış mı (1-2 zirve), çok yönlü mü (3-4 yüksek), dengeli mi? En güçlü 3 zekanın sinerjisini açıkla. En zayıf alanların akademik etkisini tartış. Minimum 2 paragraf.)*

## 2.3 Potansiyel ↔ Performans Dengesi
*(Zeka ve yetenek puanları ile çalışma davranışı ve kaygı skorları arasındaki boşluğu analiz et. Bu öğrenci potansiyelinin yüzde kaçını kullanıyor? Potansiyel kaybının nedenleri neler? Her iddia puanla kanıtlanmalı. Minimum 3 paragraf.)*

## 2.4 Çalışma Davranışı Derinlikli Analiz
*(Varsa: 7 alt kategorinin (A-G) her birini ayrı ayrı yorumla, birbirleriyle ilişkilendir. Motivasyon yüksek ama zaman yönetimi düşükse → neden? Not alma güçlü ama sınava hazırlık zayıfsa → neden? Minimum 2 paragraf.)*

---

# ⚡ 3. ÇELİŞKİ ve PARADOKS ANALİZİ

*(Bu bölüm raporun en değerli kısmıdır. Veriler arasındaki ÇELİŞKİLERİ, UYUMSUZLUKLARI ve PARADOKSLARI tespit et. Her çelişki için 3 katmanlı analiz yap:)*

| # | Çelişki Tanımı | Test 1 (Puan) | Test 2 (Puan) | Olası Açıklama | Müdahale Önerisi |
|---|---------------|---------------|---------------|----------------|-----------------|
| 1 | ... | ... | ... | ... | ... |
| 2 | ... | ... | ... | ... | ... |
| 3 | ... | ... | ... | ... | ... |
| 4 | ... | ... | ... | ... | ... |

*(Minimum 4 çelişki bul. Her biri için ayrıntılı paragraf açıklaması yaz.)*

---

# 📊 4. KAPSAMLI DEĞERLENDIRME MATRİSİ

## 4.1 Güç Envanteri

| # | Güçlü Alan | Kaynak Test | Puan | Akademik Yansıma | Sosyal Yansıma | Kariyer Potansiyeli |
|---|-----------|-------------|------|-------------------|----------------|-------------------|
| 1 | ... | ... | ... | ... | ... | ... |
| 2 | ... | ... | ... | ... | ... | ... |
| 3 | ... | ... | ... | ... | ... | ... |
| 4 | ... | ... | ... | ... | ... | ... |
| 5 | ... | ... | ... | ... | ... | ... |

*(Minimum 5 güçlü alan. Her biri farklı testlerden veya çapraz korelasyonlardan gelmeli.)*

## 4.2 Gelişim Alanları Analizi

| # | Gelişim Alanı | Kaynak Test | Puan | Risk Düzeyi | Neden Önemli? | Somut Müdahale Stratejisi |
|---|-------------|-------------|------|-------------|---------------|--------------------------|
| 1 | ... | ... | ... | 🔴/🟡 | ... | ... |
| 2 | ... | ... | ... | ... | ... | ... |
| 3 | ... | ... | ... | ... | ... | ... |
| 4 | ... | ... | ... | ... | ... | ... |

*(Minimum 4 gelişim alanı.)*

## 4.3 Kritik Göstergeler Paneli

### 🟢 Güçlü Düzey — Sürdürülmesi Gereken Alanlar
*(Puanlarla listele. Neden sürdürülmeli, nasıl daha ileri taşınabilir?)*

### 🟡 Takip Gerektiren — Potansiyel Risk Alanları
*(Puanlarla listele. Şu an kritik değil ama ihmal edilirse ne olur?)*

### 🔴 Acil İlgi — Öncelikli Müdahale Alanları
*(Puanlarla listele. Neden acil? Müdahale edilmezse 6 ay sonra ne olur?)*

---

# 🗺️ 5. STRATEJİK YOL HARİTASI

## 5.1 Akademik Başarı Planı

### 📐 Ders Bazlı Öğrenme Stratejileri

| Ders | Öğrenme Stili Uyumu | Önerilen Yöntem | Araç/Materyal | Günlük Süre |
|------|---------------------|-----------------|---------------|-------------|
| Matematik | ... | ... | ... | ... dk |
| Fen Bilimleri | ... | ... | ... | ... dk |
| Türkçe/Edebiyat | ... | ... | ... | ... dk |
| Sosyal Bilimler | ... | ... | ... | ... dk |
| Yabancı Dil | ... | ... | ... | ... dk |

*(Her dersin stratejisi VARK stili + Beyin dominansı + Çoklu Zeka profilinden türetilmeli.)*

### 📅 Haftalık Çalışma Programı Taslağı
*(Öğrencinin veri profiline özel — kaygı yüksekse kısa bloklar, motivasyon yüksekse yoğun periyotlar, kinestetik baskınsa hareket araları vb. Gün gün, saat saat örnek program.)*

### 📝 Sınav Hazırlık Protokolü
*(Sınav Kaygısı alt boyutlarına özel:)*
- **Sınavdan 1 hafta önce:** ...
- **Sınavdan 1 gün önce:** ...
- **Sınav sabahı:** ...
- **Sınav anında:** ...
- **Sınav sonrasında:** ...

## 5.2 Kişisel Gelişim Planı

### Duygusal Düzenleme Stratejileri
*(Kişilik tipi ve kaygı profiline özel. Genel "nefes al" tavsiyesi değil; bu öğrencinin spesifik stres tetikleyicilerine yönelik somut teknikler.)*

### Sosyal Beceri Geliştirme
*(Kişilik profiline göre: çok sosyalse sınır koyma, içe dönükse güvenli ortam stratejileri, çatışmacıysa empati geliştirme vb.)*

### Motivasyon ve Hedef Yönetimi
*(Kişilik tipinin motivasyon kaynaklarına uygun hedef koyma ve takip sistemi. Somut araçlar öner.)*

## 5.3 Kariyer Ön Değerlendirme Raporu

### Kariyer Yönelim Üçgeni
*(Holland RIASEC + Çoklu Zeka + Kişilik profili sentezi)*

**3 Harfli Holland Kodu Analizi:** *(Kodun ne anlama geldiği, hangi iş ortamlarında mutlu olacağı)*

**Kariyer Haritası:**

| # | Meslek / Alan | RIASEC Uyumu | Zeka Uyumu | Kişilik Uyumu | Uyum Skoru |
|---|-------------|-------------|------------|---------------|-----------|
| 1 | ... | ... | ... | ... | ⭐⭐⭐⭐⭐ |
| 2 | ... | ... | ... | ... | ⭐⭐⭐⭐⭐ |
| 3 | ... | ... | ... | ... | ⭐⭐⭐⭐ |
| 4 | ... | ... | ... | ... | ⭐⭐⭐⭐ |
| 5 | ... | ... | ... | ... | ⭐⭐⭐⭐ |
| 6 | ... | ... | ... | ... | ⭐⭐⭐ |
| 7 | ... | ... | ... | ... | ⭐⭐⭐ |
| 8 | ... | ... | ... | ... | ⭐⭐⭐ |
| 9 | ... | ... | ... | ... | ⭐⭐ |
| 10 | ... | ... | ... | ... | ⭐⭐ |

**Lise Alan Seçimi Tavsiyesi:** *(Sayısal / Eşit Ağırlık / Sözel / Dil — gerekçesiyle)*

**Üniversite Bölüm Önerileri:** *(En uygun 5 bölüm ve neden)*

**Kariyer Keşif Adımları:** *(Staj, gönüllülük, iş gölgeleme, kulüp, online kurs önerileri)*

⚠️ *Not: Bu değerlendirme bir kesin yönlendirme değil, veri destekli ön analizdir. Kesin kararlar profesyonel kariyer danışmanlığı ile desteklenmelidir.*

---

# 👨‍👩‍👦 6. AİLE DANIŞMANLIK REHBERİ

## Bu Çocuğu Anlamak

*(Ebeveynin çocuğunu daha iyi anlamasını sağlayacak, teknik terim kullanmadan yazılmış 2-3 paragraf. "Çocuğunuz şu tip bir insan..." tonunda, sıcak ve açıklayıcı.)*

## ✅ EVDEKİ DESTEK STRATEJİLERİ (Yapınız)

1. ... *(Kişilik tipine özel — neden bu yaklaşım?)*
2. ... *(Öğrenme stiline özel — somut örnek)*
3. ... *(Kaygı profiline özel — sınav döneminde nasıl davranılmalı?)*
4. ... *(Motivasyon yapısına özel — ödül/ceza dengesi)*
5. ... *(Sosyal gelişim için — arkadaşlık, aktivite önerileri)*

## ❌ KAÇINILMASI GEREKEN YAKLAŞIMLAR (Yapmayınız)

1. ... *(Kişilik tipine göre hangi baskı türü zarar verir?)*
2. ... *(Bu çocukla hangi iletişim tarzı ters etki yapar?)*
3. ... *(Hangi karşılaştırmalar motivasyonu öldürür?)*
4. ... *(Hangi beklentiler gerçekçi değil?)*

## 🗣️ EBEVEYN İLETİŞİM REHBERİ

*(Bu kişilik tipindeki bir çocukla konuşurken kullanılabilecek örnek cümleler:)*
- Başarı durumunda: "..."
- Başarısızlık durumunda: "..."
- Motivasyon düştüğünde: "..."
- Çatışma anında: "..."

---

# 👩‍🏫 7. ÖĞRETMEN VE REHBER ÖĞRETMEN REHBERİ

## Sınıf İçi Stratejiler
*(Bu öğrenci için sınıf ortamında yapılabilecek 5 somut adım. Her biri öğrenme stili ve kişilik verisine dayalı.)*

## İletişim Rehberi
*(Bu öğrenciyle en etkili iletişim tarzı. Hangi geri bildirim yöntemi işe yarar? Hangi yaklaşımlardan kaçınılmalı?)*

## Erken Uyarı İşaretleri
*(Dikkat edilmesi gereken davranış değişiklikleri — bu profildeki bir öğrencide hangi işaretler stres/tükenmişlik göstergesi olabilir?)*

## Rehber Öğretmen İçin Not
*(Bireysel görüşmelerde odaklanılması gereken temalar, izlenmesi gereken gelişim alanları)*

---

# 📌 8. SONUÇ ve ÖNCELİK MATRİSİ

## Eylem Öncelik Sıralaması

| Öncelik | Alan | Aciliyet | Sorumlu | Beklenen Süre | Başarı Göstergesi |
|---------|------|----------|---------|---------------|-------------------|
| 1. 🔴 ACİL | ... | Bu hafta | ... | ... | ... |
| 2. 🔴 ACİL | ... | 2 hafta | ... | ... | ... |
| 3. 🟡 ÖNCELİKLİ | ... | 1 ay | ... | ... | ... |
| 4. 🟡 ÖNCELİKLİ | ... | 1 ay | ... | ... | ... |
| 5. 🟢 UZUN VADE | ... | 3 ay | ... | ... | ... |
| 6. 🟢 UZUN VADE | ... | 6 ay | ... | ... | ... |

## Takip Önerisi
*(Ne zaman yeniden değerlendirme yapılmalı? Hangi alanlar 3 ay sonra tekrar ölçülmeli?)*

## Kapanış Notu
*(3-4 cümlelik profesyonel, umut verici ve güçlendirici kapanış. Bu öğrencinin en parlak potansiyelini vurgula.)*

---

*Bu rapor, EĞİTİM CHECK UP psikometrik değerlendirme sistemi tarafından, yapay zeka destekli derinlikli analiz altyapısıyla üretilmiştir. Raporda yer alan tüm yorumlar, öğrencinin psikometrik test verilerine dayanmaktadır. Bu rapor klinik tanı içermez ve klinik değerlendirme yerine geçmez.*

*Dil: Türkçe. Üslup: Profesyonel, sıcak, yapıcı, dengeli ve gerçekçi. Rapor boyunca öğrenciyi asla yargılama — potansiyelini veriye dayalı olarak ortaya koy. Abartılı övgü veya aşırı iyimser ifadelerden kaçın; samimi ve bilimsel temelli bir dil kullan. Bilimsel terimler kullandığında parantez içinde yalın Türkçe açıklamasını ekle.*"""


def _get_test_specific_guidance(test_name):
    """Her test için ticari kalite özel analiz yönergesi döndürür."""

    if "Enneagram" in test_name:
        return """
## 🔬 ENNEAGRAM KİŞİLİK TESTİ — UZMAN ANALİZ PROTOKOLÜ

Bu test, 9 Enneagram kişilik tipini %0-100 ölçeğinde ölçmektedir. Raporda aşağıdaki ANALİZ KATMANLARININ HER BİRİNİ eksiksiz ve derinlikli şekilde ele al:

### KATMAN 1: ANA TİP DERİN PROFİLİ
- Ana tipin ismi, temel motivasyonu, temel korkusu ve temel arzusu
- Bu tipin "dünya görüşü" — hayata hangi pencereden bakıyor?
- Sağlıklı düzey (büyüme modunda) → ortalama düzey → sağlıksız düzey (stres modunda) arasında bu öğrenci nerede duruyor? Puan yüzdesine göre değerlendir
- Bu tipin okul ortamındaki tipik davranış kalıpları:
  → Sınıfta nasıl oturur, nasıl dinler, nasıl katılır?
  → Ödevlere yaklaşımı nasıldır?
  → Sınav döneminde nasıl davranır?
  → Grup çalışmasında hangi rolü üstlenir?
  → Öğretmenle ilişkisi nasıldır?
  → Akranlarla ilişkisi nasıldır?
  → Başarı ve başarısızlık karşısında nasıl tepki verir?
- Bu tipin öğrenme tarzını ve akademik motivasyon kaynaklarını ayrıntıla

### KATMAN 2: KANAT (WING) ANALİZİ
- Ana tipin yanındaki iki tipten (kanat adayları) hangisinin puanı daha yüksek?
- Tam kanat notasyonu (örn: "4w5", "7w8") ve bu kombinasyonun anlamı
- Kanat etkisinin kişiliğe kattığı nüanslar (Örn: 4w3 dışa dönük ve hırslıyken, 4w5 içe dönük ve analitiktir)
- Kanat etkisinin öğrenme stili ve akademik motivasyon üzerindeki somut yansıması
- Diğer kanattan gelen zayıf etki de varsa bunu not et

### KATMAN 3: TRİTYPE (ÜÇ MERKEZ) ANALİZİ
- **Karın Merkezi (8-9-1):** Bu merkezden en yüksek puanlı tip → İçgüdüsel tepkiler, öfke yönetimi, sınır koyma
- **Kalp Merkezi (2-3-4):** Bu merkezden en yüksek puanlı tip → Duygusal tepkiler, kimlik duygusu, ilişki ihtiyacı
- **Kafa Merkezi (5-6-7):** Bu merkezden en yüksek puanlı tip → Zihinsel tepkiler, kaygı yönetimi, bilgi işleme
- Bu üç tipin birleşiminin çizdiği bütüncül portre — "Bu öğrenci stresle karşılaşınca önce ne yapar, sonra ne hisseder, sonra nasıl düşünür?"
- Tritype kombinasyonunun akademik ve sosyal hayattaki somut yansımaları

### KATMAN 4: STRES ve BÜYÜME DİNAMİĞİ
- Ana tipin stres yönündeki tip hangisi? Bu tipin puanı nedir? (Yüksekse → stres altında bu yöne kayma eğilimi güçlü)
- Ana tipin büyüme yönündeki tip hangisi? Bu tipin puanı nedir? (Yüksekse → sağlıklı gelişim potansiyeli kuvvetli)
- Stres altında bu öğrencinin sergileyeceği SOMUT davranışlar:
  → Sınıfta nasıl değişir?
  → Arkadaş ilişkilerinde ne olur?
  → Ders çalışma alışkanlıkları nasıl bozulur?
  → Bedensel belirtiler neler olabilir?
- Büyüme yolunda ilerlerken gözlemlenmesi beklenen POZİTİF değişimler

### KATMAN 5: PUAN HARİTASI ANALİZİ (9 TİP BİRLİKTE)
- Tüm 9 tipin puanlarını yüksekten düşüğe sırala ve şeklini yorumla:
  → Tek zirve profili: Ana tip belirgin, diğerleri düşük → Net, güçlü kişilik yapısı
  → Çift zirve: İki tip yakın → İç çatışma veya zenginlik göstergesi
  → Plato profili: Birden fazla tip orta-yüksek → Esnek ama belirsiz kimlik
  → Dağ silsilesi: 3-4 tip kümeleniyor → Alt grup analizi gerekli (hangi merkezde kümeleniyor?)
- En düşük puanlı tiplerin anlamı: Baskılanan, reddedilen veya gelişmemiş yönler
- İkincil ve üçüncül güçlü tiplerin ana tiple etkileşimi (destekliyor mu, çelişiyor mu?)
- Genel puan dağılımının "kişilik esnekliği" hakkında ne söylediğini yorumla

### KATMAN 6: KİŞİSEL GELİŞİM ve REHBERLIK
- Bu kişilik tipinin büyüme yolundaki 7 somut adım (yaşa uygun, günlük hayata uygulanabilir)
- Her adım için "Bunu neden yapmalı?" açıklaması
- Bu tipin düşebileceği 5 tuzak ve her birinden nasıl kaçınılır
- Aile iletişim rehberi: Bu tipte bir çocukla konuşurken kullanılması gereken dil ve yaklaşım
- Öğretmen iletişim rehberi: Sınıf ortamında bu tipi desteklemenin en etkili yolları
- Bu tipin "süper gücü" — en iyi versiyonunda dünyaya ne katar?
"""

    elif "Çalışma Davranışı" in test_name:
        return """
## 🔬 ÇALIŞMA DAVRANIŞI ÖLÇEĞİ — UZMAN ANALİZ PROTOKOLÜ

Bu test 7 alt kategoride ders çalışma alışkanlıklarını ölçer. Raporda aşağıdaki ANALİZ KATMANLARININ HER BİRİNİ eksiksiz ve derinlikli şekilde ele al:

### KATMAN 1: 7 KATEGORİ DERİN PROFİLİ
Her kategoriyi EN AZ 1 PARAGRAF derinliğinde, bu öğrencinin spesifik puanlarına dayalı olarak analiz et:

**A — Motivasyon ve Ders Çalışmaya Karşı Tutum:**
- Bu öğrencinin ders çalışmaya bakış açısı nasıl? İçsel motivasyon mu, dışsal baskı mı?
- Puan %60 altıysa: Motivasyon kaybının olası kaynakları neler? (akademik hayal kırıklığı, aile baskısı, akran etkisi, anlam bulamama)
- Puan %60 üstüyse: Bu motivasyon sürdürülebilir mi? Hangi koşullarda kırılabilir?
- Motivasyonun diğer 6 kategoriyi nasıl etkilediğini somut örneklerle göster

**B — Zaman Yönetimi:**
- Bu öğrenci zamanı nasıl kullanıyor? Planlı mı, reaktif mi, kaotik mi?
- Düşük puan: Zamanın "nereye gittiği" analizi — sosyal medya, erteleme, öncelik belirleyememe
- Yüksek puan: Planlama alışkanlığı oturmuş mu yoksa aşırı kontrol ihtiyacı mı var?
- Zaman yönetiminin sınav başarısıyla doğrudan ilişkisini bu öğrenci özelinde açıkla

**C — Derse Hazırlık ve Katılım:**
- Derse gelirken hazır mı? Önceki konuyu gözden geçiriyor mu?
- Sınıf içi katılım düzeyi: Soru soruyor mu, tartışmaya katılıyor mu, pasif mi?
- Hazırlık eksikliğinin "kartopu etkisi"ni açıkla: Bir dersi kaçırmak → sonrakini anlamamak → motivasyon kaybı

**D — Okuma ve Not Tutma Alışkanlıkları:**
- Aktif okuma mu yapıyor yoksa pasif göz gezdirme mi?
- Not tutma stratejisi: Aynen yazma, anahtar kelime, şema, zihin haritası, hiç tutmama?
- Bu öğrencinin okuma-anlama hızı ve derinliği hakkında puandan çıkarılabilecek ipuçları

**E — Yazılı Anlatım ve Ödev Yapma:**
- Ödevlere yaklaşımı: Zamanında mı, son dakika mı, kaliteli mi, geçiştirme mi?
- Yazılı ifade becerisi: Düşüncelerini organize edebiliyor mu?
- Ödev kalitesi ile not ortalaması arasındaki ilişki

**F — Sınava Hazırlanma:**
- Sınav hazırlık stratejisi: Düzenli tekrar mı, son gece maratonu mu, ezber mi, anlayarak mı?
- Sınav takvimi kullanıyor mu? Konuları bölebiliyor mu?
- Sınav sonrası değerlendirme yapıyor mu? (hataları analiz etme alışkanlığı)

**G — Genel Çalışma Koşulları ve Alışkanlıkları:**
- Fiziksel çalışma ortamı: Masa, ışık, sessizlik, telefon mesafesi
- Çalışma süresi ve verimlilik: Kaç saat çalışıyor ve bu sürenin ne kadarı gerçekten verimli?
- Beslenme, uyku ve egzersizin çalışma performansına etkisi

### KATMAN 2: ÇAPRAZ İLİŞKİ ve DARBOĞAZ ANALİZİ
- 7 kategoriyi birbirleriyle ilişkilendir — hangi kategoriler birbirini besliyor, hangileri engelliyor?
- Tipik kombinasyon kalıplarını tespit et:
  → "Yüksek motivasyon + düşük zaman yönetimi" = İstekli ama plansız öğrenci
  → "Yüksek hazırlık + düşük sınav performansı" = Kaygı kaynaklı blokaj
  → "Düşük motivasyon + yüksek ödev" = Dışsal baskıyla çalışan, tükenmişlik riski
  → "Yüksek zaman yönetimi + düşük okuma" = Planı var ama etkisiz yöntem kullanıyor
- **Ana darboğaz tespiti:** Hangi tek kategori, diğer tüm kategorilerin performansını aşağı çekiyor? Bu darboğazı kaldırmak neden 6 kategoriyi de yükseltir?
- **Kaldıraç analizi:** Hangi güçlü kategori, zayıf kategorileri kaldırmak için araç olarak kullanılabilir?

### KATMAN 3: ÇALIŞMA TİPİ PROFİLLEME
7 kategorinin puan desenine göre bu öğrencinin "çalışma tipi"ni belirle ve adlandır:
- **Disiplinli Plancı:** Tüm kategoriler dengeli yüksek → Sürdürülebilirlik ve tükenmişlik riski analizi
- **Motivasyonlu Kaotik:** Yüksek motivasyon, düşük organizasyon → Yapılandırma ihtiyacı
- **Sessiz Potansiyel:** Düşük motivasyon ama yüksek beceri → Anlam arayışı, ilgi alanı keşfi
- **Son Dakikacı Başarıcı:** Düşük hazırlık, yüksek sınav → Kısa vadeli hafıza kullanımı, uzun vadeli risk
- **Mükemmeliyetçi Yorgun:** Yüksek her şey ama stres belirtileri → Dengeleme ihtiyacı
- **Kaybolmuş Gezgin:** Çoğu kategori düşük → Acil yapılandırılmış destek
- Bu tipin avantajlarını ve risklerini detaylandır

### KATMAN 4: KİŞİYE ÖZEL EYLEM PLANI
**Günlük Çalışma Programı (hafta içi + hafta sonu ayrı):**
| Saat Dilimi | Hafta İçi | Hafta Sonu |
|------------|-----------|------------|
| 15:00-16:00 | ... | ... |
| 16:00-16:30 | ... | ... |
| ... devamı ... | ... | ... |
- Her saat dilimine ne yapılacağını, hangi tekniğin kullanılacağını yaz
- Mola zamanları, atıştırmalık, hareket araları dahil et

**Haftalık Plan Takvimi:**
| Gün | Ana Odak | Tekrar Konusu | Aktivite |
|----|----------|--------------|----------|
| Pazartesi-Pazar her gün için doldur |

**Sınav Dönemi Özel Protokolü:**
- Sınavdan 2 hafta önce → 1 hafta önce → 3 gün önce → sınav akşamı → sınav sabahı
- Her dönem için özel çalışma stratejisi, saat planı, beslenme ve uyku önerisi

### KATMAN 5: SOMUT ARAÇ ve TEKNİK REÇETESİ
- **Pomodoro tekniği:** Bu öğrenciye uygun süre ayarı (standart 25dk mı, 15dk mı, 40dk mı?)
- **Cornell not alma sistemi:** Nasıl uygulanır, örnek şablon
- **Feynman tekniği:** Konuyu "birine anlatır gibi" çalışma yöntemi
- **Aktif hatırlama (active recall):** Flash kart, soru-cevap, kendini test etme
- **Aralıklı tekrar (spaced repetition):** Anki/Quizlet kullanımı
- **Dijital araç önerileri:** Forest (odaklanma), Google Calendar (planlama), Notion (not), Todoist (görev) — her biri neden ve nasıl
- **Fiziksel çalışma ortamı reçetesi:** Masa düzeni, ışık açısı, sıcaklık, telefon mesafesi, su/atıştırmalık hazırlığı

### KATMAN 6: AİLE ve ÖĞRETMEN REHBERİ
- **Aileye:**
  → Bu çocuğun çalışma alışkanlığındaki en büyük engel nedir ve aile bunu nasıl destekleyebilir?
  → "Ders çalış!" demek NEDEN işe yaramıyor ve bunun yerine ne yapılmalı? (5 alternatif yaklaşım)
  → Çalışma ortamı hazırlamada ailenin rolü (somut kontrol listesi)
  → Ödül-ceza dengesini nasıl kurmalı? Bu öğrencinin motivasyon tipine göre ödül sistemi
  → Sınav döneminde ailenin yapması ve YAPMAMASI gerekenler (5'er madde)
  → Haftalık "10 dakikalık aile check-in" rutini: Ne sorulmalı, ne sorulMAMALI?

- **Öğretmene:**
  → Bu öğrenciyi sınıf içinde nasıl desteklenmeli? (3 somut strateji)
  → Ödev verirken dikkat edilmesi gereken hususlar
  → Geri bildirim verme yöntemi (bu çalışma profiline en uygun yaklaşım)
"""

    elif "Sağ-Sol Beyin" in test_name:
        return """
## 🔬 SAĞ-SOL BEYİN DOMINANSI TESTİ — UZMAN ANALİZ PROTOKOLÜ

Bu test beyin yarım küre baskınlığını ölçer (sağ/sol yüzde + baskınlık seviyesi). Raporda aşağıdaki ANALİZ KATMANLARININ HER BİRİNİ eksiksiz ve derinlikli şekilde ele al:

### KATMAN 1: DOMINANS PROFİLİ DERİN ANALİZİ
- Baskınlık derecesini sayısal farkla yorumla:
  → %50-55 fark: Hafif baskınlık — iki yarım küre arasında esnek geçiş yapabilir
  → %56-65 fark: Orta baskınlık — belirgin bir bilişsel tercih oluşmuş
  → %66-75 fark: Güçlü baskınlık — düşünme stili netleşmiş, zayıf taraf gelişime ihtiyaç duyar
  → %76+: Çok güçlü baskınlık — tek kanallı düşünme riski, karşı hemisfer atıl kalabilir
- Baskın yarım kürenin "dünya görüşü" — bu öğrenci hayatı, dersleri, ilişkileri ve sorunları nasıl algılıyor?
- Bu baskınlığın okul ortamındaki somut yansımaları:
  → Sınıfta nasıl oturur, nasıl dinler? (Örn: sol baskın öğrenci düzenli not alır, sağ baskın öğrenci şekil çizer)
  → Ödevlere yaklaşımı nasıldır? (Planlı-sıralı mı, son dakika-yaratıcı atılım mı?)
  → Sınavda nasıl davranır? (Baştan sona sıralı mı, kolay soruyu önce mi?)
  → Grup çalışmasında hangi rolü üstlenir? (Organize edici mi, fikir üretici mi?)
  → Hangi derslerde doğal olarak parlar, hangilerinde zorlanır ve NEDEN?
- Baskınlık profilinin günlük yaşam alışkanlıklarına etkisi (oda düzeni, zaman algısı, karar verme stili)

### KATMAN 2: İKİ HEMİSFER KARŞILAŞTIRMALI ANALİZ
- **Sol Hemisfer Profili (bu öğrencide %X):**
  → Analitik düşünme: Parçadan bütüne, adım adım, mantıksal sıralama
  → Dil ve sözel beceri: Kelime hazinesi, gramer, yazılı ifade
  → Matematiksel-sayısal işlem: Formül, hesaplama, sıralı problem çözme
  → Zaman yönetimi: Planlama, takvim, sıralı görev tamamlama
  → Detay odaklılık: Küçük parçaları fark etme, hata yakalama
  → Bu öğrencinin sol hemisfer puanının günlük yaşamdaki 3 somut gözlenebilir davranışı

- **Sağ Hemisfer Profili (bu öğrencide %X):**
  → Bütüncül düşünme: Büyük resmi görme, bağlam, ilişkilendirme
  → Görsel-mekansal beceri: Harita, şekil, resim, 3 boyutlu düşünme
  → Yaratıcılık ve sezgi: Farklı çözümler üretme, "aha!" anları, beyin fırtınası
  → Duygusal zeka: Yüz ifadesi okuma, empati, sosyal ipuçlarını yakalama
  → Müzikal ve ritmik beceri: Melodi, ritim, ses tonu farkındalığı
  → Bu öğrencinin sağ hemisfer puanının günlük yaşamdaki 3 somut gözlenebilir davranışı

- İki hemisfer arasındaki DENGEYİ veya DENGESİZLİĞİ somut örneklerle açıkla
- "Corpus callosum köprüsü" metaforu: İki yarım küre arasındaki iletişim ne kadar güçlü?

### KATMAN 3: DERS BAZLI STRATEJİ HARİTASI
Her ana ders için aşağıdaki tabloyu DOLDUR ve her satırı en az 2 cümle ile açıkla:

| Ders | Bu Öğrencinin Doğal Avantajı | Zorluk Alanı | Baskın Hemisferle Çalışma Tekniği | Zayıf Hemisferi Devreye Sokan Teknik |
|------|---------------------------|-------------|--------------------------------|-------------------------------------|
| Matematik | ... | ... | ... | ... |
| Fen Bilimleri | ... | ... | ... | ... |
| Türkçe / Edebiyat | ... | ... | ... | ... |
| Sosyal Bilimler / Tarih | ... | ... | ... | ... |
| Yabancı Dil | ... | ... | ... | ... |
| Görsel Sanatlar / Müzik | ... | ... | ... | ... |

Her ders için "çalışma senaryosu" oluştur: "Bu öğrenci Tarih konusu çalışırken önce ... yapmalı, ardından ... kullanmalı, son olarak ... ile pekiştirmeli."

### KATMAN 4: BİLİŞSEL GELİŞİM PROGRAMI
- **Baskın hemisferi sürdürme:** Bu güçlü tarafı korumak ve ileri taşımak için 5 aktivite
- **Zayıf hemisferi güçlendirme:** Atıl kalan tarafı aktive etmek için 7 günlük egzersiz programı:
  → Sol hemisfer zayıfsa: Sudoku, bulmaca, günlük yazma, liste yapma, adım adım tarif takip etme
  → Sağ hemisfer zayıfsa: Serbest çizim, müzik dinleyerek çalışma, zihin haritası, hikaye yazma, doğa yürüyüşü
- Her egzersizi "neden işe yarar" ile birlikte açıkla (nöroplastisite prensibi)
- Haftalık gelişim takvimi oluştur: Hangi gün hangi egzersiz, kaç dakika?
- 3 ay sonra beklenen gelişimsel değişimler

### KATMAN 5: KARİYER ve GELECEK YÖNELİMİ
- Bu beyin profilinin doğal olarak yatkın olduğu meslek aileleri (en az 10 meslek):
  → Sol baskın: Mühendislik, hukuk, muhasebe, yazılım, tıp, akademisyenlik, editörlük...
  → Sağ baskın: Tasarım, mimarlık, sanat, müzik, girişimcilik, pazarlama, psikoloji...
  → Dengeli: Yöneticilik, öğretmenlik, danışmanlık, araştırmacılık...
- Bu öğrencinin spesifik puan dağılımına göre "en uygun 5 kariyer yolu" ve her birinin gerekçesi
- Lise alan seçimi tavsiyesi (sayısal/eşit ağırlık/sözel) — baskınlıkla ilişkilendirerek
- Bu profil tipi üniversite sınavına nasıl hazırlanmalı? (Çalışma stratejisi)

### KATMAN 6: AİLE VE ÖĞRETMEN REHBERİ
- **Aile iletişim rehberi:**
  → Bu beyin tipiyle konuşurken kullanılması gereken dil ve yaklaşım
  → Ders çalıştırırken yapılması ve YAPILMAMASI gerekenler (en az 5'er madde)
  → Ev ortamının bu beyin tipine göre düzenlenmesi (çalışma masası, ışık, ses, renk)
  → "Bu çocuğun motivasyon anahtarı nedir?" sorusunun cevabı
- **Öğretmen rehberi:**
  → Sınıf içi bu öğrenciyi desteklemenin en etkili 5 yolu
  → Sınav/ödev esnekliği için öneriler
  → Bu öğrencinin "sıkıldığı an" nasıl fark edilir ve nasıl geri kazanılır?
  → Grup çalışmasında bu öğrenciye verilecek ideal rol
- **Bu beyin profilinin "süper gücü":** En iyi versiyonunda bu öğrenci ne başarabilir?
"""

    elif "Sınav Kaygısı" in test_name:
        return """
## 🔬 SINAV KAYGISI ÖLÇEĞİ — UZMAN ANALİZ PROTOKOLÜ

Bu test 7 alt boyutta sınav kaygısını ölçer. Raporda aşağıdaki ANALİZ KATMANLARININ HER BİRİNİ eksiksiz ve derinlikli şekilde ele al:

### KATMAN 1: 7 ALT BOYUT DERİN PROFİLİ
Her alt boyutu EN AZ 1 PARAGRAF derinliğinde, bu öğrencinin spesifik puanlarına dayalı olarak analiz et:

**1. Başkalarının Görüşü Kaygısı (Sosyal Değerlendirme):**
- Bu öğrenci sınav sonuçlarını başkalarının gözünden mi değerlendiriyor?
- Kimin görüşü en çok etkiliyor? (Aile, öğretmen, akran, toplum)
- Yüksek puan: "Ya başaramazsam ne derler?" korkusu — performans kaygısı vs değerlilik kaygısı ayrımı
- Düşük puan: Sağlıklı bir bağımsızlık mı, yoksa umursamazlık mı?
- Bu kaygının sınav anındaki somut davranışsal yansıması (kopya bakma dürtüsü, soruyu bildiği halde değiştirme vb.)

**2. Kendi Hakkındaki Görüşü (Öz-Yeterlik Algısı):**
- Bu öğrenci kendi akademik kapasitesini nasıl değerlendiriyor?
- "Ben yapamam" inancı ne kadar yerleşmiş? Kanıtları neler? (geçmiş deneyimler, karşılaştırma)
- Öz-yeterlik ile gerçek performans arasındaki uçurum analizi
- "Öğrenilmiş çaresizlik" belirtileri var mı?
- Bu boyutun motivasyon ve ders çalışma isteği üzerindeki doğrudan etkisi

**3. Gelecek Endişesi (Belirsizlik İntoleransı):**
- "Ya üniversiteyi kazanamazsam?" "Ya hayatım mahvolursa?" düzeyinde felaketleştirme var mı?
- Kaygının zaman ufku: Sadece yakın sınav mı, yoksa uzun vadeli gelecek kaygısı mı?
- Belirsizliğe tahammülsüzlük düzeyi — "kesin bilmem lazım" ihtiyacı
- Bu kaygının ders seçimi, kariyer planlaması gibi kararları nasıl etkilediği

**4. Hazırlık Endişesi (Yeterlilik Kaygısı):**
- "Yeterince çalışmadım" duygusu: Gerçekçi mi, yoksa mükemmeliyetçilik mi?
- Ne kadar çalışırsa çalışsın "yetmez" hissi var mı? (sonsuz döngü analizi)
- Hazırlık endişesi ile erteleme davranışı arasındaki paradoksal ilişki
- Bu boyutun çalışma süresini artırıp verimliliği düşürme mekanizması

**5. Bedensel Tepkiler (Somatik Kaygı):**
- Hangi bedensel belirtiler ön planda? Mide bulantısı, terleme, çarpıntı, baş ağrısı, uyuyamama, titreme, nefes darlığı
- Bu belirtiler ne zaman başlıyor? (günler önce mi, sınav sabahı mı, sınav anında mı?)
- Bedensel belirtilerin "ikincil kaygı" yaratması: "Ellerim titriyor → Yazamayacağım → Panik"
- Somatizasyon düzeyi: Kaygı bedene ne kadar yansıyor?

**6. Zihinsel Tepkiler (Bilişsel Kaygı):**
- Zihin boşalması, unutma, konsantrasyon kaybı, "bildiklerimin hepsini unuttum" deneyimi
- Ruminasyon (takıntılı düşünme): Sınavdan önce/sonra aynı düşünceleri tekrar tekrar yaşama
- "White-out" fenomeni: Sınav kağıdını görünce her şeyi unutma
- Bilişsel kapasitenin kaygı tarafından "işgal edilmesi" mekanizması (çalışma belleği yükü)
- Bu boyutun gerçek bilgi düzeyi ile sınav performansı arasındaki makası açması

**7. Genel Kaygı (Yaygın Kaygı Düzeyi):**
- Kaygı sadece sınavlara mı özgü, yoksa genel bir yaşam tutumu mu?
- Trait kaygı (kişilik özelliği) vs State kaygı (durumsal) ayrımı
- Genel kaygının diğer 6 alt boyutu nasıl beslediği ve büyüttüğü

### KATMAN 2: KAYGI PROFİLİ TİPLEME
7 alt boyutun puan desenine göre bu öğrencinin kaygı profilini belirle:
- **Bedensel Ağırlıklı Kaygı:** Somatik belirtiler dominant → Gevşeme ve beden farkındalığı öncelikli
- **Bilişsel Ağırlıklı Kaygı:** Zihinsel belirtiler dominant → Bilişsel yeniden yapılandırma öncelikli
- **Sosyal Kaynaklı Kaygı:** Başkalarının görüşü dominant → Öz-değer çalışması öncelikli
- **Varoluşsal Kaygı:** Gelecek endişesi dominant → Anlam ve perspektif çalışması öncelikli
- **Mükemmeliyetçi Kaygı:** Hazırlık endişesi dominant → "Yeterince iyi" kavramı üzerine çalışma
- **Karma Kaygı:** Birden fazla boyut yüksek → Çok boyutlu müdahale programı

### KATMAN 3: KAYGI DÖNGÜSÜ ve MEKANİZMA ANALİZİ
Bu öğrencinin spesifik puan profiline göre kişiselleştirilmiş bir kaygı döngüsü çiz:
```
Tetikleyici (sınav tarihi açıklanması)
    ↓
Otomatik Düşünce ("Bu sınavı geçemeyeceğim çünkü...")
    ↓
Duygusal Tepki (korku, endişe, panik, çaresizlik)
    ↓
Bedensel Tepki (mide, terleme, çarpıntı, uyuyamama)
    ↓
Davranışsal Tepki (erteleme / aşırı çalışma / kaçınma)
    ↓
Sonuç (düşük performans / tükenmişlik)
    ↓
Doğrulama ("İşte, yine başaramadım") → Döngü başa döner
```
- Her adımda bu öğrenciye özel somut örnekler ver
- Döngüyü kırmak için müdahale edilebilecek EN KOLAY noktayı belirle ve neden orası olduğunu açıkla

### KATMAN 4: YERKES-DODSON PERFORMANS ANALİZİ
- Kaygı-performans ilişkisini bu öğrenci özelinde değerlendir:
  → Çok düşük kaygı: Motivasyon eksikliği riski
  → Optimal kaygı: Performansı artıran sağlıklı gerginlik
  → Aşırı kaygı: Performansı engelleyen yıkıcı stres
- Bu öğrenci şu an bu eğrinin neresinde?
- Optimal bölgeye gelmek için kaygıyı ne yöne hareket ettirmeli?
- Her ders/sınav tipi için optimal kaygı seviyesi farklı mıdır?

### KATMAN 5: 5 AŞAMALI SINAV HAZIRLIK PROTOKOLÜ
Bu öğrencinin kaygı profiline özel, adım adım sınav hazırlık programı:

**Aşama 1 — Sınavdan 2 Hafta Önce:**
- Konuları günlere bölme takvimi (somut şablon)
- Günlük çalışma rutini (saat, süre, mola)
- "Endişe saati" tekniği: Günde 15 dk ayrılmış kaygı zamanı, geri kalanında erteleme

**Aşama 2 — Sınavdan 1 Hafta Önce:**
- Aktif hatırlama ve kendini test etme stratejileri
- Eksik konuları tamamlama vs mükemmeliyetçilikten vazgeçme dengesi
- Uyku düzeni oluşturma (saat önerisi)

**Aşama 3 — Sınavdan 3 Gün Önce:**
- Yeni konu ÖĞRENMEME kuralı — sadece tekrar ve pekiştirme
- Hafif fiziksel aktivite (yürüyüş, esneme, hafif koşu)
- Dijital detoks başlangıcı (sosyal medya kısıtlama)

**Aşama 4 — Sınav Akşamı:**
- Son 2 saat: Özet tarama (detaya girmeme)
- Çanta hazırlığı, kıyafet seçimi (karar yükünü azaltma)
- Uyku öncesi gevşeme rutini (4-7-8 nefes tekniği adım adım)
- "Yarına hazırım" güçlendirme cümlesi

**Aşama 5 — Sınav Sabahı ve Sınav Anı:**
- Sabah rutini (uyandıktan sınava kadar saat saat)
- Beslenme önerisi (ne yemeli, ne yememeli)
- Sınav salonuna giriş stratejisi (erken varma, tanıdıklarla sınav konuşmama)
- **İlk 5 dakika protokolü:** Kağıdı al → 3 derin nefes → tüm soruları oku → kolay soruyu bul → başla
- Sınav anı panik stratejisi: "Beynim dondu" anında yapılacak 3 adım
- Zaman yönetimi: Soru başına süre hesaplama, takılınca geçme kuralı

### KATMAN 6: BİLİŞSEL YENİDEN YAPILANDIRMA ÇALIŞMASI
Bu öğrencinin kaygı profiline özel en az 7 bilişsel yeniden yapılandırma örneği:

| Otomatik Düşünce | Bilişsel Çarpıtma Türü | Alternatif Düşünce | Kanıt |
|-------------------|----------------------|-------------------|-------|
| "Bu sınavda kesin başarısız olacağım" | Felaketleştirme | "Daha önce de zor hissettiğim sınavlarda geçtim" | Geçmiş sınav sonuçları |
| ... en az 7 satır ... |

- Her düşünce-alternatif çiftini bu öğrencinin yaşına ve durumuna uygun yaz
- "Düşünce günlüğü" tutma alışkanlığı nasıl oluşturulur (şablon ver)

### KATMAN 7: AİLE ve ÖĞRETMEN REHBERİ
**Aileye:**
- Sınav döneminde evde nasıl bir atmosfer yaratılmalı? (5 somut kural)
- "Nasıl geçti?" sorusu NEDEN yıkıcıdır ve bunun yerine ne sorulmalı?
- Baskı yapmadan nasıl destek olunur? (Destekleyici vs Baskıcı aile davranışı karşılaştırma tablosu)
- Çocuğun kaygı belirtilerini nasıl fark eder? (erken uyarı işaretleri listesi)
- Sınav sonucu kötü geldiğinde yapılması ve YAPILMAMASI gerekenler
- Ailenin kendi kaygısının çocuğa nasıl bulaştığı ve bunu engellemek için stratejiler

**Öğretmene:**
- Sınav anında kaygılı öğrenciyi nasıl fark eder ve destekler?
- Sınav formatında yapılabilecek kaygı azaltıcı düzenlemeler
- Geri bildirim verirken kaygıyı artırmayan dil kullanımı
- Sınıf genelinde kaygı azaltıcı ortam yaratma önerileri

**⚠️ NOT:** Klinik düzeyde kaygı belirtileri gözlemlenirse, bu değerlendirme profesyonel psikolojik destek önerisini destekler.
"""

    elif "VARK" in test_name:
        return """
## 🔬 VARK ÖĞRENME STİLLERİ TESTİ — UZMAN ANALİZ PROTOKOLÜ

Bu test 4 öğrenme kanalını ölçer: V (Görsel), A (İşitsel), R (Okuma/Yazma), K (Kinestetik). Raporda aşağıdaki ANALİZ KATMANLARININ HER BİRİNİ eksiksiz ve derinlikli şekilde ele al:

### KATMAN 1: 4 ÖĞRENME KANALI DERİN PROFİLİ
Her kanalı EN AZ 1 PARAGRAF derinliğinde, bu öğrencinin spesifik puanlarına dayalı olarak analiz et:

**V — Görsel (Visual) Kanal (Bu öğrencide: ?):**
- Bilgiyi görsel imgeler, diyagramlar, renkler, şekiller ve mekansal düzenlemelerle işler
- Bu öğrencide bu kanal ne kadar aktif? Puan düzeyi ne anlama geliyor?
- Günlük yaşamda gözlemlenebilir belirtiler: Tahtaya mı bakıyor, şekil çiziyor mu, renkli kalemler kullanıyor mu, yüz tanıma becerisi güçlü mü?
- Bu kanalın güçlü olması: Harita okuma, grafik yorumlama, şema oluşturma, zihin haritası becerisi
- Bu kanalın zayıf olması: Sözel yönergelerde kaybolma, tahtadaki şekilleri anlamada zorluk

**A — İşitsel (Aural) Kanal (Bu öğrencide: ?):**
- Bilgiyi dinleyerek, tartışarak, sesli düşünerek ve ritmik tekrarlarla işler
- Günlük yaşam belirtileri: Dersi dinleyerek mi öğreniyor, sesli çalışıyor mu, müzik dinlerken çalışabiliyor mu, sözlü talimatları kolayca takip ediyor mu?
- Bu kanalın güçlü olması: Tartışma grupları, podcast/sesli kitap, öğretmenin anlatımından çok şey alma
- Bu kanalın zayıf olması: Uzun açıklamalarda dikkat dağılması, sesli ortamda çalışamama

**R — Okuma/Yazma (Read/Write) Kanal (Bu öğrencide: ?):**
- Bilgiyi yazılı metinler, listeler, tanımlar, notlar ve yazma eylemiyle işler
- Günlük yaşam belirtileri: Kitap okumayı seviyor mu, detaylı not tutuyor mu, yazarak çalışıyor mu, liste yapıyor mu?
- Bu kanalın güçlü olması: Ders kitabından öğrenme, özet çıkarma, essay yazma, araştırma yapma
- Bu kanalın zayıf olması: Uzun metinlerden sıkılma, not tutmada isteksizlik

**K — Kinestetik (Kinesthetic) Kanal (Bu öğrencide: ?):**
- Bilgiyi deneyimleyerek, yaparak, dokunarak, hareket ederek ve pratik uygulama ile işler
- Günlük yaşam belirtileri: Yerinde duramıyor mu, elleriyle oynuyor mu, laboratuvar ve pratiği seviyor mu, spor/el işi becerisi güçlü mü?
- Bu kanalın güçlü olması: Deney yapma, model oluşturma, alan gezisi, rol yapma, simülasyon
- Bu kanalın zayıf olması: Uzun süre oturup dinlemede zorluk, soyut kavramları somutlaştırma ihtiyacı

### KATMAN 2: ÖĞRENME MODALİTESİ ANALİZİ
4 kanalın puan dağılımına göre bu öğrencinin öğrenme modalitesini belirle:
- **Tek Baskın (Unimodal):** Bir kanal belirgin yüksek → Bu kanalı maksimize et, diğerlerini destek olarak kullan
- **Çift Baskın (Bimodal):** İki kanal yakın yüksek → İki kanalın sinerjisini açıkla (Örn: VA=Anlatım dinlerken şekil çizen, VK=Görerek ve yaparak öğrenen)
- **Üçlü Baskın (Trimodal):** Üç kanal yakın → Esnek öğrenme kapasitesi, farklı ortamlara uyum
- **Çoklu (Multimodal):** Dört kanal dengeli → Evrensel öğrenici, her yöntemden faydalanır ama hiçbirinde uzmanlaşmamış olabilir
- Bu modalite tipinin avantajları ve potansiyel riskleri
- Puan farkları analizi: En güçlü ile en zayıf kanal arasındaki fark ne anlama geliyor?

### KATMAN 3: DERS BAZLI ÖĞRENME REÇETESİ
Her ana ders için bu öğrencinin baskın stiline uygun detaylı çalışma senaryosu oluştur:

| Ders | Baskın Stile Uygun Teknik | Somut Araç/Materyal | Adım Adım Çalışma Senaryosu |
|------|--------------------------|--------------------|-----------------------------|
| **Matematik** | ... | ... | "Önce ... yap, sonra ... kullan, ardından ... ile pekiştir" |
| **Fen Bilimleri** | ... | ... | "..." |
| **Türkçe / Edebiyat** | ... | ... | "..." |
| **Sosyal Bilimler / Tarih** | ... | ... | "..." |
| **Yabancı Dil (İngilizce)** | ... | ... | "..." |
| **Görsel Sanatlar / Müzik** | ... | ... | "..." |

Her senaryo EN AZ 3 cümle olmalı ve somut, uygulanabilir adımlar içermeli.
Senaryo formatı: "Bu öğrenci [ders] konusu çalışırken: 1) ... yapmalı, 2) ... kullanmalı, 3) ... ile pekiştirmeli."

### KATMAN 4: SINAV ve HAFIZA STRATEJİLERİ
Bu öğrencinin baskın stiline özel sınav hazırlık ve hafıza teknikleri:

**Kodlama (bilgiyi kaydetme) aşaması:**
- V: Zihin haritası, renk kodlama, şema, infografik oluşturma teknikleri (adım adım)
- A: Sesli not kaydı, kendine anlatma, çalışma arkadaşıyla tartışma, müzikle eşleştirme
- R: Cornell notu, özet çıkarma, flash kart yazma, kendi sınav sorusu oluşturma
- K: Yürürken tekrar, yazı tahtasına yazma, jest ve mimiklerle eşleştirme, model yapma
→ Bu öğrencinin baskın stiline göre ÖNCELİKLİ 3 kodlama tekniğini seç ve adım adım açıkla

**Depolama (bilgiyi saklama) aşaması:**
- Aralıklı tekrar takvimi (bu stile uygun: görsel kart mı, sesli tekrar mı, yazılı test mi?)
- Uyku öncesi 10 dakika tekrarı — stil bazlı format önerisi
- Hafıza sarayı tekniğinin bu stile uyarlanması

**Geri çağırma (sınavda hatırlama) aşaması:**
- Sınav anında stil bazlı hatırlama ipuçları:
  → V: "O şemayı gözünde canlandır", renk ve konum ipuçları
  → A: "O anlatımı kulaklarında duyar gibi ol", ritmik tekrar ipuçları
  → R: "O notu yazdığın sayfayı hatırla", liste ve başlık ipuçları
  → K: "O deneyi yaparken ne hissettiğini hatırla", hareket ve dokunma ipuçları

### KATMAN 5: ZAYIF KANALLARI GÜÇLENDİRME PROGRAMI
- Neden önemli: Tek kanala bağımlılık riski ve çoklu kanal avantajı
- Her zayıf kanal için 5 günlük egzersiz programı:
  → Egzersiz adı, süresi, nasıl yapılır, neden işe yarar
- Zayıf kanalı güçlendirmenin baskın kanalı ZAYIFLATMAYACAĞI garantisi (nöroplastisite)
- 1 ay sonra beklenen gelişim ve kontrol yöntemi
- Evde yapılabilecek kanal geliştirme aktiviteleri (yaşa uygun)

### KATMAN 6: DİJİTAL ARAÇ ve KAYNAK REÇETESİ
Bu öğrencinin baskın stiline göre kişiselleştirilmiş dijital araç önerileri:
- **V için:** Canva, MindMeister, YouTube eğitim kanalları (somut kanal isimleri), infografik araçları
- **A için:** Podcast önerileri, sesli kitap uygulamaları, kayıt araçları, diksiyon uygulamaları
- **R için:** Notion, Obsidian, Anki, blog yazma platformları, e-kitap kaynakları
- **K için:** Simülasyon uygulamaları, deney videoları, maker/DIY projeleri, sportif aktiviteler
→ Bu öğrencinin baskın stiline göre "İLK 3 İNDİR" listesi oluştur

### KATMAN 7: AİLE ve ÖĞRETMEN REHBERİ
**Aileye:**
- Bu öğrencinin öğrenme stili evde nasıl desteklenmeli?
- Çalışma ortamının stile göre düzenlenmesi (masa, ışık, ses, renk, materyal)
- "Ders çalıştırırken" bu stile uygun yaklaşım (Örn: K baskın çocuğa "Otur, oku!" demek İŞE YARAMAZ)
- Aile bireylerinin kendi öğrenme stilleri farklıysa oluşan çatışma ve çözümü
- Bu stilin günlük yaşamda nasıl destekleneceği (hobi, aktivite, sohbet biçimi)

**Öğretmene:**
- Sınıf içi bu öğrenciyi desteklemenin en etkili 5 yolu
- Ödev formatını bu stile uygun alternatif şekillerle sunma önerileri
- Değerlendirme yöntemlerinde stil bazlı esneklik önerileri
- Bu stilin sınıf dinamiğine katkısı ve risk alanları
"""

    elif "Çoklu Zeka" in test_name:
        return """
## 🔬 ÇOKLU ZEKA TESTİ (GARDNER) — UZMAN ANALİZ PROTOKOLÜ

Bu test Gardner'ın 8 zeka alanını %0-100 ölçeğinde ölçer. Raporda aşağıdaki ANALİZ KATMANLARININ HER BİRİNİ eksiksiz ve derinlikli şekilde ele al:

### KATMAN 1: 8 ZEKA ALANI DERİN PROFİLİ
Her zeka alanını EN AZ 1 PARAGRAF derinliğinde, bu öğrencinin spesifik puanlarına dayalı olarak analiz et:

**1. Sözel-Dilsel Zeka (Bu öğrencide: %?):**
- Kelime hazinesi, dil kullanımı, okuma/yazma kapasitesi, ikna becerisi, hikaye anlatma
- Bu puanın okul ortamındaki yansıması: Türkçe/Edebiyat, yabancı dil derslerindeki performans
- Günlük yaşam göstergeleri: Kitap okuma isteği, tartışma becerisi, kelime oyunları, yazma alışkanlığı
- Bu düzeyin gelişim potansiyeli ve somut güçlendirme aktiviteleri

**2. Mantıksal-Matematiksel Zeka (Bu öğrencide: %?):**
- Sayısal akıl yürütme, problem çözme, örüntü tanıma, neden-sonuç ilişkisi, soyut düşünme
- Okul yansıması: Matematik, fen, kodlama derslerindeki doğal eğilim
- Günlük yaşam: Strateji oyunları, bulmaca merakı, "neden?" sorusu sıklığı, sistematik düşünme
- Gelişim potansiyeli ve aktiviteler

**3. Görsel-Uzamsal Zeka (Bu öğrencide: %?):**
- 3 boyutlu düşünme, mekansal ilişkiler, görsel hafıza, renk/şekil duyarlılığı, harita okuma
- Okul yansıması: Geometri, resim, harita, grafik yorumlama, deney şeması çizme
- Günlük yaşam: Lego/puzzle ilgisi, yön bulma becerisi, çizim/boyama merakı, fotoğraf çekme
- Gelişim potansiyeli ve aktiviteler

**4. Bedensel-Kinestetik Zeka (Bu öğrencide: %?):**
- Beden kontrolü, el becerisi, koordinasyon, dokunsal öğrenme, fiziksel ifade
- Okul yansıması: Beden eğitimi, laboratuvar, el işi, drama, dans, spor
- Günlük yaşam: Sportif yetenek, el işi becerisi, hareket ihtiyacı, tamir/yapma merakı
- Gelişim potansiyeli ve aktiviteler

**5. Müzikal-Ritmik Zeka (Bu öğrencide: %?):**
- Melodi algılama, ritim duygusu, ses tonu farkındalığı, müzikal bellek, enstrüman yeteneği
- Okul yansıması: Müzik dersi, ritmik sayma, şiir ezberinde ritim kullanma
- Günlük yaşam: Şarkı söyleme, enstrüman çalma, müzik dinleme tercihleri, ses taklidi
- Gelişim potansiyeli ve aktiviteler

**6. Kişilerarası (Sosyal) Zeka (Bu öğrencide: %?):**
- Empati, liderlik, iş birliği, çatışma çözme, ikna, sosyal ipuçlarını okuma
- Okul yansıması: Grup çalışması, sınıf başkanlığı, akran arabuluculuğu, takım sporları
- Günlük yaşam: Arkadaşlık kalitesi, sosyal çevre genişliği, organize etme isteği
- Gelişim potansiyeli ve aktiviteler

**7. İçsel (Özedönük) Zeka (Bu öğrencide: %?):**
- Öz-farkındalık, duygusal okuryazarlık, kendi güçlü/zayıf yanlarını bilme, içsel motivasyon
- Okul yansıması: Bağımsız çalışma tercihi, öz-değerlendirme becerisi, kendi hedeflerini koyma
- Günlük yaşam: Günlük tutma, yalnız vakit geçirme ihtiyacı, derin düşünme, felsefe ilgisi
- Gelişim potansiyeli ve aktiviteler

**8. Doğacı Zeka (Bu öğrencide: %?):**
- Doğa gözlemi, sınıflandırma, çevre duyarlılığı, hayvan/bitki ilgisi, ekolojik düşünme
- Okul yansıması: Fen/biyoloji, çevre projeleri, arazi çalışmaları, koleksiyon yapma
- Günlük yaşam: Doğa yürüyüşü, hayvan sevgisi, bahçe işleri, doğa belgeselleri
- Gelişim potansiyeli ve aktiviteler

### KATMAN 2: ZEKA PROFİLİ ŞEKİL ANALİZİ
8 zekanın puan dağılımının şeklini yorumla:
- **Uzmanlaşmış Profil (1-2 zirve):** Belirgin güçlü alanlar → Derinleşme stratejisi
- **Çok Yönlü Profil (3-4 yüksek):** Geniş yetenek yelpazesi → Çapraz alan sinerjisi
- **Dengeli Profil (hepsi orta):** Esnek ama zirvesiz → Potansiyel keşif ihtiyacı
- **Dağınık Profil (yüksek-düşük karışık):** Belirgin güçlü ve zayıf yanlar → Stratejik gelişim
- Bu öğrencinin profil tipi hangisi ve bunun anlamı nedir?

**"Zeka İmzası" Tespiti:**
- En güçlü 3 zekanın birleşimini tek bir cümlede tanımla: "Bu öğrenci, [X] zekasının [özelliği] ile [Y] zekasının [özelliği]ni birleştiren, [Z] zekasıyla desteklenen bir [etiket] profilidir."
- Bu imzanın dünyada nasıl bir iz bırakabileceğini anlat

**Sinerjik Zeka Etkileşimleri:**
- Hangi güçlü zekalar birbirini destekliyor? (Örn: Sözel + Sosyal = Doğal lider/konuşmacı)
- Hangi zeka çiftleri çelişiyor? (Örn: Yüksek İçsel + Düşük Sosyal = Yalnız çalışmayı tercih)
- Bu etkileşimlerin okul yaşamındaki somut yansımaları

### KATMAN 3: ZEKA-DERS EŞLEŞTİRME HARİTASI
Kapsamlı bir tablo oluştur:

| Zeka Alanı | Puan | Doğal Olarak Güçlü Olduğu Dersler | Bu Zekayı Kullanan Çalışma Tekniği | Evde Güçlendirme Aktivitesi |
|-----------|------|--------------------------------|-----------------------------------|---------------------------|
| Her 8 zeka için doldur — her satır en az 2 cümle açıklama içermeli |

**Bu öğrencinin güçlü zekalarını kullanarak zayıf dersleri nasıl öğrenebileceğini somut senaryolarla göster:**
- Örn: Müzikal zeka güçlü + Matematik zayıf → "Formülleri ritmik tekerleme yaparak ezberle, çarpım tablosunu şarkı haline getir"
- Örn: Kinestetik zeka güçlü + Tarih zayıf → "Tarihi olayları canlandırarak öğren, harita üzerinde yürüyerek ülkeleri keşfet"
- Bu öğrenciye özel en az 5 çapraz strateji geliştir

### KATMAN 4: ZEKA-KARİYER EŞLEŞTİRME HARİTASI
Bu öğrencinin güçlü zeka kombinasyonuna göre kariyer yönlendirmesi:

| Güçlü Zeka Kombinasyonu | Uygun Kariyer Alanları | Ünlü Rol Modeller | Bu Öğrenci İçin Somut İlk Adım |
|------------------------|---------------------|-------------------|-------------------------------|
| Her güçlü zeka çifti/üçlüsü için en az 5 meslek önerisi |

- Kariyer önerilerini Türkiye iş piyasası gerçekleriyle ilişkilendir
- Lise alan seçimi tavsiyesi (bu zeka profiline göre)
- Üniversite bölüm önerileri (en az 8 bölüm, her birinin gerekçesiyle)
- "21. yüzyıl meslekleri" arasında bu profile uygun yeni nesil kariyer yolları

### KATMAN 5: KİŞİSEL GELİŞİM PROGRAMI
**Güçlü zekaları derinleştirme (ilk 3):**
- Her güçlü zeka için: kulüp, yarışma, kurs, atölye, proje önerileri
- "Uzmanlık yolu" haritası: Bugün → 6 ay sonra → 1 yıl sonra → 3 yıl sonra

**Zayıf zekaları güçlendirme (son 2):**
- Her zayıf zeka için haftalık 3 egzersiz (süre, yöntem, beklenen etki)
- Neden zayıf zekaları da geliştirmek gerekir? (nöroplastisite, bütüncül gelişim)

**Ders dışı aktivite programı:**
- Bu zeka profiline en uygun 7 ders dışı aktivite (her biri için neden ve nasıl)
- Yaz tatili gelişim planı
- Hafta sonu aktivite önerileri

### KATMAN 6: AİLE ve ÖĞRETMEN REHBERİ
**Aileye:**
- Bu çocuğun en güçlü zekaları üzerinden nasıl motive edilir?
- "Bu çocuk neden [şunu] seviyor ama [bunu] yapmak istemiyor?" sorusunun zeka profiliyle açıklaması
- Ev ortamında zeka geliştirme: Her oda bir öğrenme fırsatı (mutfak=matematik, bahçe=doğacı, salon=sosyal...)
- Hobi ve kurs seçiminde zeka profilini rehber olarak kullanma
- Bu profildeki bir çocukla iletişim kurmanın en etkili yolu
- Yapılması ve YAPILMAMASI gerekenler (en az 5'er madde)

**Öğretmene:**
- Bu öğrencinin sınıf içi "parlama anları" ne zaman ve nasıl oluşur?
- Farklılaştırılmış öğretim stratejileri (bu zeka profiline özel ödev/proje alternatifleri)
- Değerlendirmede zeka dostu yaklaşımlar (sadece yazılı sınav yerine alternatif değerlendirme)
- Bu öğrenciye "doğru görevi" vermek: Grup çalışmasında, sınıf etkinliğinde, projede ideal rol
- Bu öğrencinin "süper gücü" — en iyi versiyonunda sınıfa ve dünyaya ne katar?
"""

    elif "Holland" in test_name:
        return """
## 🔬 HOLLAND MESLEKİ İLGİ ENVANTERİ (RIASEC) — UZMAN ANALİZ PROTOKOLÜ

Bu test 6 mesleki ilgi tipini 0-28 puan aralığında ölçer. Raporda aşağıdaki ANALİZ KATMANLARININ HER BİRİNİ eksiksiz ve derinlikli şekilde ele al:

### KATMAN 1: 6 TİP DERİN PROFİLİ
Her tipi EN AZ 1 PARAGRAF derinliğinde, bu öğrencinin spesifik puanlarına dayalı olarak analiz et:

**R — Gerçekçi (Realistic) Tip (Bu öğrencide: ?/28):**
- Somut, pratik, fiziksel, el becerisi, araç-gereç, doğa, yapı-inşa, tamir
- Bu puanın anlamı: Yüksekse → somut sonuçlar üreten işler ister, düşükse → masabaşı/teorik işleri tercih eder
- Okul yansıması: Atölye, laboratuvar, teknoloji-tasarım, beden eğitimi, tarım, el sanatları
- Bu öğrencide R tipinin günlük yaşamdaki 3 somut göstergesi
- R tipi insanların çalışma ortamı tercihi: Açık hava, atölye, fabrika, laboratuvar, inşaat sahası

**I — Araştırmacı (Investigative) Tip (Bu öğrencide: ?/28):**
- Merak, analiz, araştırma, sorgulama, keşfetme, bilimsel yöntem, veri, hipotez
- Okul yansıması: Fen, matematik, bilim olimpiyatları, araştırma projeleri, deney tasarlama
- Bu öğrencide I tipinin günlük yaşamdaki göstergeleri (soru sorma sıklığı, okuma tercihleri, belgesel izleme)
- I tipi insanların çalışma ortamı: Üniversite, araştırma merkezi, laboratuvar, kütüphane

**A — Sanatçı (Artistic) Tip (Bu öğrencide: ?/28):**
- Yaratıcılık, özgünlük, estetik, hayal gücü, ifade özgürlüğü, kural dışılık
- Okul yansıması: Görsel sanatlar, müzik, drama, edebiyat, yaratıcı yazarlık, tasarım
- Bu öğrencide A tipinin göstergeleri (giyim tarzı, oda dekorasyonu, hobi seçimi, müzik tercihi)
- A tipi insanların çalışma ortamı: Stüdyo, sahne, atölye, serbest çalışma, yaratıcı ajans

**S — Sosyal (Social) Tip (Bu öğrencide: ?/28):**
- Yardım etme, öğretme, iyileştirme, rehberlik, empati, ekip çalışması, iletişim
- Okul yansıması: Grup çalışmaları, akran danışmanlığı, gönüllülük, liderlik, sosyal projeler
- Bu öğrencide S tipinin göstergeleri (arkadaş çevresi, yardım etme isteği, dinleme becerisi)
- S tipi insanların çalışma ortamı: Okul, hastane, danışma merkezi, STK, sosyal hizmet

**E — Girişimci (Enterprising) Tip (Bu öğrencide: ?/28):**
- Liderlik, ikna, risk alma, rekabet, yönetme, satış, strateji, etki yaratma
- Okul yansıması: Sınıf/okul başkanlığı, münazara, proje liderliği, kulüp kurma, organizasyon
- Bu öğrencide E tipinin göstergeleri (karar verme tarzı, ikna becerisi, para/iş merakı)
- E tipi insanların çalışma ortamı: Ofis, toplantı odası, sahne, pazar, yönetim katı

**C — Geleneksel (Conventional) Tip (Bu öğrencide: ?/28):**
- Düzen, detay, doğruluk, sistem, prosedür, veri organizasyonu, güvenilirlik
- Okul yansıması: Düzenli not tutma, zamanında ödev teslimi, kural takibi, arşivleme
- Bu öğrencide C tipinin göstergeleri (masa düzeni, planlama alışkanlığı, detaycılık)
- C tipi insanların çalışma ortamı: Ofis, muhasebe, kütüphane, banka, devlet dairesi

### KATMAN 2: 3 HARFLİ HOLLAND KODU DERİN ANALİZİ
- En yüksek 3 tipi belirle → Holland kodu (Örn: "IAS", "SEC", "RIA")
- **Kod yorumlama:** Bu üç harfin birleşik anlamını detaylı açıkla — tek tek değil, BİRLİKTE ne söylüyorlar?
- **Holland altıgeni (hexagon) analizi:**
  → Bitişik tipler (Örn: RI, IA, AS, SE, EC, CR) = Uyumlu, doğal geçiş
  → Karşıt tipler (Örn: R-S, I-E, A-C) = Gerilim veya zenginlik göstergesi
  → Bu öğrencinin kodundaki tiplerin altıgendeki konumları uyumlu mu, çelişkili mi?
- **Kodun "hikayesi":** Bu kod, nasıl bir iş ortamı, yaşam tarzı ve değer sistemi arayan birini tarif ediyor?
- **Kodun puan güvenilirliği:** İlk 3 tip arasındaki puan farkları analizi — net ayrışma var mı yoksa belirsiz mi?

### KATMAN 3: KAPSAMLI KARİYER HARİTASI
Bu öğrencinin 3 harfli koduna göre detaylı meslek önerileri:

| # | Meslek / Alan | Holland Kodu Uyumu | Gerekli Eğitim | Türkiye'de İş İmkanı | Maaş Potansiyeli | Bu Öğrenci İçin Neden Uygun? |
|---|-------------|-------------------|---------------|---------------------|-----------------|---------------------------|
| 1-20 arası doldur — EN AZ 20 meslek önerisi |

Meslekleri aşağıdaki kategorilere ayırarak sun:
- **💎 En Uygun 5 Meslek:** Tüm 3 tip ile yüksek uyum
- **🔵 İyi Uyumlu 5 Meslek:** 2 tip ile uyum
- **🟢 Keşfedilmeye Değer 5 Meslek:** 1 tip ile uyum ama diğer özelliklere de hitap eden
- **🆕 Gelecek Meslekleri 5 Meslek:** Yapay zeka, dijital dönüşüm, yeşil ekonomi alanlarından

### KATMAN 4: EĞİTİM YÖNLENDİRME
**Lise Alan Seçimi:**
- Sayısal / Eşit Ağırlık / Sözel / Dil → hangisi ve NEDEN (Holland koduyla bağlantılı gerekçe)
- Bu alanda başarılı olmak için gereken özellikler ve bu öğrencinin profiliyle eşleştirme
- "Ya yanlış alan seçersem?" endişesini ele alma

**Üniversite Bölüm Önerileri:**
| # | Bölüm | Üniversite Önerileri (Türkiye) | Holland Uyumu | Gelecek Vizyonu |
|---|-------|-------------------------------|-------------|----------------|
| 1-10 arası en az 10 bölüm önerisi |

**Alternatif Eğitim Yolları:**
- Yurt dışı eğitim düşünülüyorsa: Uygun ülkeler ve bölümler
- Meslek yüksekokulu alternatifi (bu profile uygunsa)
- Sertifika programları ve online eğitim yolları

### KATMAN 5: KARİYER KEŞİF ve EYLEM PLANI
Bu öğrencinin kariyer keşif yolculuğu için somut adımlar:

**Hemen Yapılabilecek (Bu Hafta):**
- Online kariyer testi çapraz doğrulaması (KPSS meslek tercihi, İŞKUR mesleki eğilim vb.)
- Holland koduna uygun YouTube kanalları, belgeseller, kitap önerileri (somut isimlerle)
- Sosyal medyada takip edilecek meslek profesyonelleri

**Kısa Vadeli (1-3 Ay):**
- İş gölgeleme (job shadowing): Bu meslekleri yakından görmek için kimlerle iletişime geçilmeli?
- Staj/gönüllülük fırsatları (yaşa uygun)
- Kariyer fuarları, üniversite tanıtım günleri

**Orta Vadeli (3-12 Ay):**
- Kulüp, yarışma, proje katılımları (Holland koduna uygun)
- Yaz okulu/kampı önerileri
- Mentorluk programları

**Uzun Vadeli (1-3 Yıl):**
- Portfolyo/CV oluşturma alışkanlığı
- Yetkinlik geliştirme planı (bu kariyer alanı için hangi beceriler gerekli?)
- Üniversite başvuru stratejisi

### KATMAN 6: AİLE ve ÖĞRETMEN REHBERİ
**Aileye:**
- Bu çocuğun mesleki ilgi profili ne söylüyor? (Aile dilinde, jargonsuz açıklama)
- "Ben doktor/avukat/mühendis olmasını istiyordum" durumunda: Aile beklentisi vs çocuğun doğal eğilimi çatışmasını nasıl yönetmeli?
- Bu profildeki bir çocuğu motive etmenin en etkili yolu (hangi ödüller, hangi teşvikler?)
- Kariyer keşif sürecinde ailenin rolü: Destekleyici vs yönlendirici dengesini nasıl kurar?
- Yapılması ve YAPILMAMASI gereken 5'er madde
- Bu profil tipiyle "yemek masasında kariyer sohbeti" nasıl yapılır?

**Öğretmene:**
- Bu öğrenciyi motive eden ders içi aktivite türleri (Holland tipine göre)
- Proje ödevi ve performans görevi önerirken bu profili nasıl kullanabilir?
- Kariyer günü / mesleki tanıtım etkinliklerinde bu öğrenci için özel yönlendirme
- Rehberlik servisiyle paylaşılabilecek özet bilgiler

**⚠️ YASAL UYARI:** Bu değerlendirme profesyonel kariyer danışmanlığını destekler; tek başına kesin mesleki yönlendirme için yeterli değildir. Nihai kararlar çoklu veri kaynağı ve uzman görüşüyle verilmelidir.
"""



    elif "P2 Dikkat" in test_name:
        return """
## 🔬 P2 DİKKAT TESTİ — UZMAN ANALİZ PROTOKOLÜ

Bu test, seçici dikkat ve konsantrasyon ölçen dijital bir testtir (15 satır × 30 sembol, "p" harfi hedefli, nokta bazlı). Raporda aşağıdaki ANALİZ KATMANLARININ HER BİRİNİ eksiksiz ve derinlikli şekilde ele al:

### KATMAN 1: TEMEL METRİKLER DERİN ANALİZİ
Her metriği ayrı ayrı, bu öğrencinin spesifik değerlerine dayalı olarak yorumla:

**CP (Konsantrasyon Performansı) — En Kritik Gösterge:**
- Doğru hedef isabetleri − Yanlış işaretlemeler
- Yüzdelik değerlendirme ve yaş grubuna göre yorumlama

**TN-E (Toplam Performans):** Hız ve doğruluğu birlikte değerlendiren bileşik gösterge

**E1 (Atlama Hatası):** Kaçırılan hedef sayısı — dikkat dağılması göstergesi
- Hangi satırlarda yoğunlaştığını analiz et

**E2 (Yanlış İşaretleme):** Dürtüsellik göstergesi — kontrolsüz tepki

**E1/E2 Oranı:** Hata profili tespiti (dikkat eksikliği mi, dürtüsellik mi)

**FR (Dalgalanma):** Dikkat sürdürülebilirliği

### KATMAN 2: SATIR BAZLI PERFORMANS EĞRİSİ
15 satırlık performans eğrisini detaylı analiz et:
- İlk 5 satır (ısınma): Performans nasıl başlıyor?
- Orta 5 satır (sürdürme): Stabil mi, dalgalı mı?
- Son 5 satır (yorulma): Düşüş var mı?
- Satır bazlı hata dağılımı — kaçırma ve yanlış işaretleme hangi satırlarda artıyor?

### KATMAN 3: HIZ-DOĞRULUK DENGESİ
Profil tiplemesi: Dengeli / Dürtüsel / Temkinli / Gelişen

### KATMAN 4: AKADEMİK YANSIMA VE ÖNERİLER
- Sınıf içi dikkat performansı tahmini
- Sınav stratejisi önerileri
- Dikkat geliştirme programı
- Aile ve öğretmen için spesifik tavsiyeler
- Günlük yaşamda dikkat yönetimi stratejileri

**⚠️ ÖNEMLİ NOT:** Bu test klinik bir tanı aracı DEĞİLDİR. Sonuçlar dikkat eğilimlerini gösterir.
"""

    elif "Akademik Analiz" in test_name:
        return """
## 🔬 AKADEMİK ANALİZ TESTİ — UZMAN ANALİZ PROTOKOLÜ

Bu test, 4 alt boyutta akademik yetkinliği ölçen performans bazlı bir testtir (İlköğretim: 44 soru, Lise: 54 soru). Raporda aşağıdaki ANALİZ KATMANLARININ HER BİRİNİ eksiksiz ve derinlikli şekilde ele al:

### KATMAN 1: OKUMA ANLAMA DERİN ANALİZİ
Bu öğrencinin okuma anlama puanını (%?) derinlemesine analiz et:

**Metin Kavrama Becerisi:**
- Ana fikir yakalama: Metinde "ne anlatılıyor?" sorusuna verdiği cevapların kalitesi
- Detay hatırlama: Metindeki spesifik bilgilere ulaşabilme becerisi
- Çıkarım yapma: "Metinde doğrudan söylenmeyen ama ima edilen nedir?" sorusuna cevap verme kapasitesi
- Kelime anlamı bağlamdan çıkarma: Bilinmeyen kelimelerin anlamını metinden çıkarabilme

**Bu Puanın Akademik Yansımaları:**
- Okuma anlama TÜM derslerin temelidir — matematik problemi bile önce okuyarak anlaşılır
- Bu puanla öğrenci, ders kitabından bağımsız öğrenebilir mi?
- Uzun metinleri (sınav soruları, araştırma makaleleri, kitap) ne kadar etkili işleyebilir?
- Metin tabanlı sınavlarda (LGS paragraf soruları, YKS Türkçe) beklenen performans

**Gelişim Stratejileri:**
- %80+: İleri düzey okuma materyalleri, eleştirel okuma, çapraz metin karşılaştırma
- %60-79: Aktif okuma stratejileri (SQ3R yöntemi), günlük 20 dk okuma rutini
- %40-59: Yapılandırılmış okuma programı, metin türlerine göre strateji
- <%40: Temel okuma becerisi güçlendirme, sesli okuma, kelime hazinesi çalışması

### KATMAN 2: MATEMATİKSEL MUHAKEME DERİN ANALİZİ
Bu öğrencinin matematiksel muhakeme puanını (%?) derinlemesine analiz et:

**Matematiksel Düşünme Kapasitesi:**
- Temel işlem yetkinliği: Dört işlem, kesir, yüzde hesaplama becerisi
- Problem çözme: Sözel problemi matematiksel modele dönüştürebilme
- Soyut düşünme: Formül, denklem, değişken kavramlarını anlama (yaşa göre)
- Çok adımlı muhakeme: Birden fazla adım gerektiren problemlerde performans
- Örüntü tanıma: Sayı dizileri, seri tamamlama, matematiksel ilişkileri görme

**Bu Puanın Akademik Yansımaları:**
- Fen derslerinde formül kullanma ve grafik yorumlama kapasitesi
- Günlük yaşam matematiksel okuryazarlığı (bütçe, oran, istatistik)
- Sınav performansı tahmini: Matematiksel soruları hızlı ve doğru çözebilme kapasitesi

**Gelişim Stratejileri:**
- %80+: Olimpiyat problemleri, ileri düzey mantık soruları, kodlama/programlama
- %60-79: Çok adımlı problem pratiği, farklı soru tiplerine maruz kalma
- %40-59: Temel kavramları somutlaştırma (manipülatifler, görsel modeller), günlük pratik
- <%40: Sayı kavramı güçlendirme, temel işlem otomatikleştirme, birebir destek

### KATMAN 3: MANTIKSAL DÜŞÜNME DERİN ANALİZİ
Bu öğrencinin mantıksal düşünme puanını (%?) derinlemesine analiz et:

**Analitik Düşünme Kapasitesi:**
- Analoji kurma: "A, B'ye oranla ne ise; C, D'ye oranla odur" düşünebilme
- Seri/dizi tamamlama: Sayı, harf, şekil dizilerinde kuralı bulma
- Sıralama ve sınıflandırma: Bilgiyi organize etme, kategorize etme becerisi
- Mantıksal çıkarım: "Eğer X ise, o zaman Y" tarzı çıkarımlar yapabilme
- Eleştirel düşünme (lise): Argüman değerlendirme, tutarsızlık tespiti, varsayım sorgulama

**Bu Puanın Yaşam Yansımaları:**
- Karar verme kalitesi: Mantıksal düşünme, sağlıklı karar vermenin temelidir
- Problem çözme yaklaşımı: Sorunlara sistematik mi, rastgele mi yaklaşıyor?
- Akademik başarı ile korelasyonu: Mantıksal düşünme, tüm derslerdeki başarıyı etkiler
- Bu öğrencinin "akıl yürütme tarzı": Tümevarımsal mı (örneklerden kurala) yoksa tümdengelimsel mi (kuraldan örneğe)?

**Gelişim Stratejileri:**
- %80+: Felsefe tartışmaları, münazara, kodlama, strateji oyunları (satranç, Go)
- %60-79: Mantık bulmacaları, Sudoku, analoji çalışmaları, "neden?" sorusu alışkanlığı
- %40-59: Temel mantık egzersizleri, somut-soyut geçiş çalışmaları, sınıflandırma oyunları
- <%40: Somut manipülatiflerle mantık kurma, basit "eğer-ise" çalışmaları, örüntü oyunları

### KATMAN 4: PERFORMANS vs ÖZ-DEĞERLENDİRME UYUM ANALİZİ
Bu öğrencinin öz-değerlendirmesi (%?) ile gerçek performansı (%?) arasındaki ilişkiyi derinlemesine analiz et:

**Uyum Analizi:**
- Fark ≤%10 → **Tutarlı (Sağlıklı Farkındalık):**
  → Bu öğrenci kendini gerçekçi değerlendiriyor — sağlıklı bir akademik benlik algısı var
  → Bu tutarlılığın avantajları: Gerçekçi hedef koyma, etkili çalışma planlaması
  → Risk: Yok, ama gelişim motivasyonunu canlı tutmak için yeni hedefler gerekli

- Öz-değerlendirme > Performans (+%10+) → **Aşırı Özgüven:**
  → Bu öğrenci kendini olduğundan daha başarılı algılıyor
  → Olası nedenler: Sosyal karşılaştırma eksikliği, aile beklentisinin içselleştirilmesi, gerçekçi geri bildirim alamama
  → Riskler: Yetersiz çalışma, "zaten biliyorum" yanılgısı, sınav hayal kırıklığı
  → Müdahale: Sınav analizleri ile gerçeklik testi, spesifik hedefler, "neyi bilmiyorsun?" sorusu

- Öz-değerlendirme < Performans (-%10+) → **Düşük Özgüven:**
  → Bu öğrenci aslında başarılı ama bunu görmüyor/kabul etmiyor
  → Olası nedenler: Yüksek aile beklentisi, mükemmeliyetçilik, geçmiş akademik travma, olumsuz geri bildirim geçmişi
  → Riskler: Motivasyon kaybı, "nasıl olsa yapamam" tuzağı, potansiyelin altında performans
  → Müdahale: Başarı kanıtlarını somutlaştırma, pozitif geri bildirim, "bunu SEN başardın" vurgusu

**Öz-Değerlendirme Alt Maddeleri Analizi:**
- Odaklanma algısı, tekrar çalışma alışkanlığı, sınav özgüveni, ödev disiplini, araştırma isteği
- Hangi alt maddelerde kendini yüksek, hangilerinde düşük değerlendirmiş?
- Bu alt maddelerin gerçek performans bölümleriyle çapraz karşılaştırması

### KATMAN 5: BÜTÜNLEŞİK AKADEMİK PROFİL
4 boyutun etkileşim analizi — tüm verileri bir araya getirerek kapsamlı bir profil çiz:

**Güçlü Alan → Zayıf Alan Transfer Stratejisi:**
- En güçlü alanın yöntemlerini en zayıf alana nasıl aktarılır?
- Somut örnekler: "Okuma anlama güçlü → Matematik problem metnini önce bir okuma-anlama ödevi gibi ele al"
- Çapraz beceri geliştirme planı

**Akademik Potansiyel Tahmini:**
- 4 boyutun ortalaması ve dağılımı: Dengeli mi, tek ayaklı mı?
- Bu profilin LGS/YKS gibi merkezi sınavlarda beklenen performansı
- Güçlü alanların kariyer yönelimi ipuçları

**Kişiye Özel 3 Aşamalı Gelişim Planı:**

**Aşama 1 — Acil (0-1 Ay):**
- En zayıf alandaki temel eksiklikleri kapatma
- Günlük 15 dakikalık hedefli çalışma rutini
- Hızlı kazanımlar için kolay hedefler

**Aşama 2 — Orta Vade (1-3 Ay):**
- Zayıf alanı ortalamaya taşıma
- Güçlü alanları derinleştirme
- Haftalık ilerleme ölçümü

**Aşama 3 — Uzun Vade (3-6 Ay):**
- Tüm alanları dengeli yüksek seviyeye taşıma
- İleri düzey materyal ve yarışma hazırlığı
- Kontrol testi ile ilerleme değerlendirmesi

### KATMAN 6: AİLE ve ÖĞRETMEN REHBERİ
**Aileye:**
- Bu çocuğun akademik profili ne söylüyor? (Aile dilinde, jargonsuz, 3 cümlede özet)
- En güçlü alan üzerinden motivasyon kurma: "Bu alanda çok iyisin, bu yeteneğini [zayıf alan] için de kullanabilirsin"
- Ev ortamında akademik destek: Hangi alanda nasıl yardım edilebilir?
- Özel ders / destek eğitimi gerekiyor mu? Hangi alanda, ne düzeyde?
- Ders çalışma sırasında yapılması ve YAPILMAMASI gerekenler
- Sınav sonuçlarına nasıl tepki verilmeli? (başarılı ve başarısız sonuç için ayrı strateji)
- Ekran süresi, oyun ve sosyal medyanın akademik performansa etkisi ve yönetimi

**Öğretmene:**
- Bu öğrencinin sınıf içi akademik profili: Nerede parlar, nerede zorlanır?
- Farklılaştırılmış öğretim: Bu öğrenciye uygun zorluk seviyesi ve materyal
- Grup çalışmasında bu öğrenciye verilecek ideal rol (güçlü alanına göre)
- Geri bildirim verirken dikkat edilecekler (özgüven durumuna göre yaklaşım)
- Rehberlik servisine yönlendirme gerekiyor mu? Hangi konuda?
"""

    return ""


def build_single_test_prompt(student_name, student_age, student_gender, test_name, test_data, student_grade=None):
    """Tekil test analizi için ticari kalite prompt — her teste özel uzman protokolü içerir."""

    test_guidance = _get_test_specific_guidance(test_name)
    grade_text = f"{student_grade}. Sınıf" if student_grade else "Belirtilmemiş"

    return f"""# ROL ve KİMLİK

Sen, Türkiye'nin önde gelen eğitim psikolojisi merkezlerinde 20 yıl deneyim kazanmış, psikometrik test yorumlama konusunda uzmanlaşmış bir Klinik Eğitim Psikoloğusun.

Bu rapor, ücretli bir profesyonel danışmanlık hizmetinin çıktısıdır. Tek bir test sonucunu, sanki karşında oturan aileye yüz yüze sunuyormuş gibi, derinlikli, kişiselleştirilmiş ve uygulanabilir şekilde analiz edeceksin.

---

# ÖĞRENCİ DOSYASI

| Alan | Bilgi |
|------|-------|
| İsim | {student_name} |
| Yaş | {student_age} |
| Sınıf | {grade_text} |
| Cinsiyet | {student_gender} |
| Analiz Edilen Test | {test_name} |
| Değerlendirme Türü | Tekil Test Derinlikli Analiz |

## TEST VERİSİ (JSON)
```json
{json.dumps(test_data, ensure_ascii=False, indent=2)}
```

---

# KRİTİK KURALLAR

1. **KANITSAL ZORUNLULUK:** Her yorum, iddia ve tespit mutlaka parantez içinde test adı ve puan ile desteklenmeli. Kanıtsız hiçbir yorum yapma.

2. **DERİNLİK ZORUNLULUĞU:** Bu ücretli bir profesyonel hizmettir. Her bölüm, bir psikolog danışmanlık seansında anlatacağı kadar detaylı olmalı. Genel geçer, şablonik, "daha çok çalış" tarzı yüzeysel tavsiyeler YASAK. Her öneri bu öğrencinin spesifik puan profilinden türetilmeli.

3. **PUAN YORUMLAMA ÇERÇEVESİ:**
   - %0-20 → Belirgin gelişim alanı — yapılandırılmış destek önerilir
   - %21-40 → Ortalamanın altı — hedefli çalışma gerektirir
   - %41-60 → Ortalama düzey — strateji ile yükseltilebilir
   - %61-80 → Güçlü alan — sürdürülebilir ve ileri taşınabilir
   - %81-100 → Çok güçlü — yetenek göstergesi, özel destekle parlayabilir

4. **TIBBİ TANI YASAĞI:** Klinik tanı terimleri (DEHB, depresyon, disleksi, anksiyete bozukluğu vb.) kesinlikle kullanma.

5. **GELİŞİMSEL BAĞLAM:** {student_age} yaşında, {grade_text} düzeyinde bir bireyin gelişimsel özelliklerini referans al.

6. **UZUNLUK:** Bu rapor minimum 2500 kelime olmalıdır. Her bölüm ödenen ücrete değecek derinlikte olmalı.

7. **BİLİMSEL TEMEL ZORUNLULUĞU:** Tüm yorumlar bilimsel araştırmalara dayalı olmalıdır. Bilimsel terim veya kavram kullandığında hemen ardından parantez içinde lise öğrencisinin bile anlayabileceği yalın bir Türkçe açıklama ekle. Örnek: "Metakognitif farkındalık (kendi düşünme sürecini izleyebilme ve yönetebilme becerisi)..." veya "Yerkes-Dodson Yasası'na göre (belirli bir noktaya kadar kaygı performansı artırır ama aşılırsa performans düşer)..."

8. **DENGELİ ve GERÇEKÇİ MOTİVASYON TONU:** Abartılı motivasyon ifadelerinden kaçın. "Muhteşem", "olağanüstü", "inanılmaz potansiyel", "her şeyi başarabilir" gibi abartılı ifadeler YASAK. Gerçekçi, dengeli ve yapıcı bir ton kullan. Güçlü yönleri veriye dayalı belirt, gelişim alanlarını görmezden gelme. "Bu alanda güçlü bir profil ortaya koyuyor" gibi ölçülü ifadeler tercih et. Sakin, profesyonel bir danışman üslubu — ne karamsar ne de abartılı iyimser.

9. **YALIN TÜRKÇE ZORUNLULUĞU:** Rapor, okuryazar bir lise öğrencisinin rahatlıkla anlayabileceği düzeyde yazılmalı. Teknik terimler kullanıldığında mutlaka parantez içinde kısa, günlük dile yakın açıklama ekle. Kısa ve net cümleler kur. Yabancı kökenli sözcükler yerine mümkünse Türkçe karşılıklarını tercih et.

---
{test_guidance}
---

# RAPOR FORMATI (HER BÖLÜMÜ AYNEN DOLDUR, HİÇBİRİNİ ATLAMA)

---

## 📋 YÖNETİCİ ÖZETİ
*(Raporu okuyacak kişinin 1 dakikada tüm tabloyu göreceği 4-5 cümlelik güçlü özet. En kritik bulgu, en önemli güç, en acil gelişim alanı ve en öncelikli adım.)*

---

## 📊 1. TEST SONUÇ TABLOSU

**Tek Cümle Sonuç:** *(Testin en önemli bulgusunu, öğrenciyi tanımayan birinin bile anlayacağı netlikte ifade et.)*

**Tüm Boyutlar Görsel Özeti:**
```
[Boyut/Kategori Adı]    : ██████████ XX%  → [Kısa Yorum]
[Boyut/Kategori Adı]    : ████████░░ XX%  → [Kısa Yorum]
[Boyut/Kategori Adı]    : ██████░░░░ XX%  → [Kısa Yorum]
[Boyut/Kategori Adı]    : ████░░░░░░ XX%  → [Kısa Yorum]
...devam — TÜM boyutları listele, hiçbirini atlama
```

---

## 🧠 2. DERİNLEMESİNE YORUM

*(Bu raporun kalbi burasıdır. Her alt boyutu/kategoriyi ayrı ayrı derinlemesine yorumla ve birbirleriyle ilişkilendir.

Her alt boyut için:
- Bu puan ne anlama geliyor?
- Günlük hayatta nasıl gözlemlenir?
- Okul ortamında nasıl yansır?
- Diğer alt boyutlarla nasıl etkileşir?

Ardından genel profil sentezi:
- Profilin şekli — dengeli mi, tek zirve mi, çoklu zirve mi?
- Bu profilin "hikayesi" — veriler birlikte okunduğunda ne anlatıyor?
- Öğrencinin bu profile sahip olmasının olası gelişimsel ve çevresel nedenleri

Minimum 5-6 paragraf, akıcı ve profesyonel anlatım.)*

---

## 💪 3. GÜÇLÜ YÖNLER ANALİZİ

| # | Güçlü Yön | Kanıt (Puan) | Okul Yaşamında Nasıl Gözlemlenir? | Nasıl İleri Taşınabilir? | Kariyer Bağlantısı |
|---|-----------|--------------|----------------------------------|--------------------------|-------------------|
| 1 | ... | ... | ... | ... | ... |
| 2 | ... | ... | ... | ... | ... |
| 3 | ... | ... | ... | ... | ... |
| 4 | ... | ... | ... | ... | ... |
| 5 | ... | ... | ... | ... | ... |

*(Minimum 5 güçlü yön. Her birini 1-2 cümlelik açıklamayla destekle.)*

---

## 🌱 4. GELİŞİM ALANLARI ve MÜDAHALE STRATEJİLERİ

| # | Gelişim Alanı | Mevcut Durum (Puan) | Risk Düzeyi | Bu Neden Önemli? | Haftalık Gelişim Planı |
|---|-------------|---------------------|-------------|-----------------|----------------------|
| 1 | ... | ... | 🔴/🟡/🟢 | ... | ... |
| 2 | ... | ... | ... | ... | ... |
| 3 | ... | ... | ... | ... | ... |
| 4 | ... | ... | ... | ... | ... |

*(Minimum 4 gelişim alanı. Her biri için detaylı strateji.)*

---

## 🎯 5. KAPSAMLI AKSİYON PLANI

**📌 STRATEJİ 1: [Başlık]**
- **Hedef:** *(Ne başarılacak?)*
- **Neden bu öğrenci için önemli:** *(Veri referansıyla)*
- **Adım adım uygulama:** *(Günlük/haftalık program)*
- **Gerekli araç/materyal:** *(Somut)*
- **Başarı göstergesi:** *(Nasıl ölçülecek?)*
- **Sorumlu:** *(Öğrenci/Öğretmen/Aile)*
- **Beklenen süre:** *(Ne kadar sürede sonuç görülür?)*

**📌 STRATEJİ 2: [Başlık]**
*(Aynı formatta)*

**📌 STRATEJİ 3: [Başlık]**
*(Aynı formatta)*

**📌 STRATEJİ 4: [Başlık]**
*(Aynı formatta)*

**📌 STRATEJİ 5: [Başlık]**
*(Aynı formatta)*

---

## 👨‍👩‍👦 6. AİLE DANIŞMANLIK BÖLÜMÜ

### Bu Sonuçlar Ne Anlama Geliyor?
*(Teknik terminolojiyi aile diline çevir. Ebeveynin çocuğunu daha iyi anlamasını sağla. 2-3 paragraf.)*

### ✅ Evde Yapılması Gerekenler (En Az 5 Madde)
*(Her madde test verisine dayalı, somut ve uygulanabilir. "Neden?" açıklaması ile.)*

### ❌ Kaçınılması Gerekenler (En Az 4 Madde)
*(Kişilik/profil tipine göre hangi yaklaşımlar zarar verebilir? Somut örneklerle.)*

### 🗣️ İletişim Rehberi
*(Bu profildeki bir çocukla nasıl konuşulmalı? Duruma göre örnek cümleler:)*
- Başarı gösterdiğinde: "..."
- Zorlandığında: "..."
- Motivasyonu düştüğünde: "..."
- Çatışma anında: "..."

---

## 👩‍🏫 7. ÖĞRETMEN ve REHBER ÖĞRETMEN BÖLÜMÜ

### Sınıf İçi Stratejiler (En Az 5 Madde)
*(Her strateji bu öğrencinin veri profilinden türetilmiş olmalı.)*

### İletişim ve Geri Bildirim Yaklaşımı
*(Bu öğrenciyle en etkili iletişim tarzı. Nelere dikkat edilmeli?)*

### Erken Uyarı İşaretleri
*(Bu profildeki bir öğrencide hangi davranış değişiklikleri risk göstergesi olabilir?)*

### Rehber Öğretmen İçin Takip Planı
*(Bireysel görüşmelerde odaklanılacak temalar, izlenecek gelişim alanları)*

---

## 📌 8. SONUÇ ve ÖNCELİK MATRİSİ

| Öncelik | Eylem | Aciliyet | Sorumlu | Süre | Başarı Göstergesi |
|---------|-------|----------|---------|------|-------------------|
| 1. 🔴 | ... | Bu hafta | ... | ... | ... |
| 2. 🔴 | ... | 2 hafta | ... | ... | ... |
| 3. 🟡 | ... | 1 ay | ... | ... | ... |
| 4. 🟡 | ... | 1 ay | ... | ... | ... |
| 5. 🟢 | ... | 3 ay | ... | ... | ... |

### Takip Önerisi
*(Ne zaman yeniden değerlendirme yapılmalı?)*

### Kapanış Notu
*(Profesyonel, umut verici, güçlendirici kapanış. Bu öğrencinin potansiyelini vurgula.)*

---

*Bu rapor, EĞİTİM CHECK UP psikometrik değerlendirme sistemi tarafından, yapay zeka destekli derinlikli analiz altyapısıyla üretilmiştir. Raporda yer alan tüm yorumlar, öğrencinin test verilerine dayanmaktadır. Bu rapor klinik tanı içermez.*

*Dil: Türkçe. Üslup: Profesyonel, sıcak, yapıcı, dengeli ve gerçekçi. Öğrenciyi asla yargılama — potansiyelini veriye dayalı olarak ortaya koy. Abartılı övgü veya aşırı iyimser ifadelerden kaçın; samimi, bilimsel temelli bir dil kullan. Bilimsel terimler kullandığında parantez içinde yalın Türkçe açıklamasını ekle.*"""

# ============================================================
# ANA ÖĞRETMEN UYGULAMASI
# ============================================================
###############################################################################
# ENTEGRE RAPORLAMA SİSTEMİ — 3'LÜ RAPOR (Öğretmen / Öğrenci / Ebeveyn)
###############################################################################

# Patronun istediği sıralama
INTEGRATED_TEST_ORDER = [
    "Öğrenme Stilleri Testi",      # 1 — Potansiyel Analiz
    "Beyin Yatkınlıkları Testi",   # 2
    "Çoklu Zeka Testi",            # 3
    "Enneagram Kişilik Testi",     # 4
    "Holland Mesleki İlgi Testi",  # 5
    "Çalışma Davranışı Ölçümü",    # 6 — Mevcut Durum Tespiti
    "Akademik Analiz",             # 7
    "Sınav Kaygısı Testi",         # 8
    "P2 Dikkat Testi",             # 9
]

# Eşleştirme: veritabanında farklı isimle saklanmış olabilir
TEST_NAME_MAP = {
    "VARK Öğrenme Stilleri Testi": "Öğrenme Stilleri Testi",
    "Vark Öğrenme Stilleri Testi": "Öğrenme Stilleri Testi",
    "Öğrenme Stilleri Testi": "Öğrenme Stilleri Testi",
    "Beyin Yatkınlıkları Testi": "Beyin Yatkınlıkları Testi",
    "Sağ-Sol Beyin Testi": "Beyin Yatkınlıkları Testi",
    "Çoklu Zeka Testi": "Çoklu Zeka Testi",
    "Çoklu Zekâ Testi": "Çoklu Zeka Testi",
    "Enneagram Kişilik Testi": "Enneagram Kişilik Testi",
    "Enneagram Testi": "Enneagram Kişilik Testi",
    "Holland Mesleki İlgi Testi": "Holland Mesleki İlgi Testi",
    "Holland Mesleki İlgi Envanteri": "Holland Mesleki İlgi Testi",
    "Çalışma Davranışı Ölçümü": "Çalışma Davranışı Ölçümü",
    "Çalışma Davranışı Ölçme Envanteri": "Çalışma Davranışı Ölçümü",
    "Akademik Analiz": "Akademik Analiz",
    "Sınav Kaygısı Testi": "Sınav Kaygısı Testi",
    "Sınav Kaygısı Değerlendirme Ölçeği": "Sınav Kaygısı Testi",
    "P2 Dikkat Testi": "P2 Dikkat Testi",
    "D2 Dikkat Testi": "P2 Dikkat Testi",
}


def _normalize_test_name(name):
    """Veritabanındaki test adını standart sıralama adına çevir."""
    return TEST_NAME_MAP.get(name, name)


def _order_tests(tests_list):
    """Testleri patronun istediği sıraya göre sıralar, bilinmeyenleri sona ekler."""
    ordered = []
    remaining = list(tests_list)
    for canonical in INTEGRATED_TEST_ORDER:
        for t in remaining[:]:
            if _normalize_test_name(t["test_name"]) == canonical:
                ordered.append(t)
                remaining.remove(t)
                break
    ordered.extend(remaining)
    return ordered


def build_integrated_report_prompt(student_name, student_age, student_gender,
                                    test_data_list, report_type="ogretmen",
                                    student_grade=None):
    """
    Patronun istediği sıralı entegre rapor promptunu oluşturur.
    report_type: 'ogretmen' | 'ogrenci' | 'ebeveyn'
    """

    audience_map = {
        "ogretmen": {
            "title": "ÖĞRETMEN / KOÇ",
            "tavsiye_hedef": "öğretmene ve koça",
            "hitap": "Öğretmene ve Koça Tavsiyeler",
            "perspective": (
                "Raporu okuyan kişi öğretmen veya eğitim koçudur. "
                "Tavsiyeleri sınıf içi uygulamalar, ders planlama, bireysel rehberlik "
                "ve akademik koçluk perspektifinden yaz."
            )
        },
        "ogrenci": {
            "title": "ÖĞRENCİ",
            "tavsiye_hedef": "öğrenciye",
            "hitap": "Öğrenciye Tavsiyeler",
            "perspective": (
                "Raporu okuyan kişi öğrencinin kendisidir. "
                "Öğrencinin yaşına uygun, samimi ve motive edici bir dil kullan. "
                "Tavsiyeleri öğrencinin kendi başına uygulayabileceği somut adımlar olarak yaz. "
                "'Sen' diye hitap et."
            )
        },
        "ebeveyn": {
            "title": "EBEVEYN",
            "tavsiye_hedef": "ebeveynlere",
            "hitap": "Ebeveynlere Tavsiyeler",
            "perspective": (
                "Raporu okuyan kişi öğrencinin anne/babasıdır. "
                "Bilimsel terimleri basit Türkçeyle açıkla. "
                "Tavsiyeleri ebeveynin evde uygulayabileceği somut adımlar olarak yaz. "
                "Çocuğun adını kullan, 'çocuğunuz' yerine ismini tercih et."
            )
        }
    }

    aud = audience_map.get(report_type, audience_map["ogretmen"])
    grade_text = f", {student_grade}. sınıf öğrencisi" if student_grade else ""

    # Test verilerini sıralı JSON olarak hazırla
    tests_json = json.dumps(test_data_list, ensure_ascii=False, indent=2)

    prompt = f"""Sen deneyimli bir eğitim psikoloğu ve psikometri uzmanısın. {student_name} ({student_age} yaşında, {student_gender}{grade_text}) adlı öğrencinin test sonuçlarını analiz edeceksin.

# RAPOR TÜRÜ: {aud['title']} RAPORU
{aud['perspective']}

# KRİTİK RAPORLAMA KURALLARI

## SIRALAMAYA UYMA ZORUNLULUĞU
Testleri aşağıdaki KESIN sıra ile raporla. Öğrenci testleri karışık sırayla çözmüş olsa bile SEN bu sırayı takip et:

**POTANSİYEL ANALİZ (İlk 5 Test):**
1. Öğrenme Stilleri Testi (VARK)
2. Sağ-Sol Beyin Yatkınlıkları Testi
   → 1 ve 2'yi birleştirip KISA bir entegre yorum yaz
3. Çoklu Zekâ Testi
   → 1, 2 ve 3'ü birleştirip KISA bir entegre yorum yaz
4. Enneagram Kişilik Testi
   → 1, 2, 3 ve 4'ü birleştirip KISA bir entegre yorum yaz
5. Holland Mesleki İlgi Envanteri
   → 5 testin tamamını "POTANSİYEL ANALİZ ÖZETİ" başlığı altında birleştir

**MEVCUT DURUM TESPİTİ (Son 4 Test):**
6. Çalışma Davranışı Ölçümü
7. Akademik Analiz
8. Sınav Kaygısı Testi
9. P2 Dikkat Testi
   → 4 testin tamamını "MEVCUT DURUM ÖZETİ" başlığı altında birleştir

**FİNAL: POTANSİYEL + MEVCUT DURUM SENTEZ**
→ Potansiyel analiz ile mevcut durumu ilişkilendirerek kapsamlı bir ÖNERİ bölümü yaz
→ Meslek yönelimi tavsiyeleri ekle

## FORMAT KURALLARI
- Her test için: Veri tablosu özeti → Kısa bilimsel yorum → {aud['hitap']}
- Birleştirme yorumları: Önceki testlerle tutarlılık/çelişki analizi, maksimum 4-5 cümle
- Potansiyel ile mevcut durumu ilişkilendirirken somut örnekler ver
- ÇOK UZATMA! Her test analizi kısa ve öz olsun. Uzun paragraflardan kaçın.
- Tavsiyeler somut, uygulanabilir ve küçük küçük olsun
- Markdown formatında yaz (başlıklar, tablolar, madde işaretleri)
- Rapor bir A4 dosya olarak düşünüldüğünde 12-15 sayfa civarında olsun

## VAROLAN TEST VERİLERİ
Aşağıdaki listede sadece öğrencinin ÇÖZMİŞ olduğu testler var. Eğer sıralamada belirtilen bir test listede yoksa o testi ATLA ve bir sonrakine geç. Atlanan test hakkında yorum yapma.

```json
{tests_json}
```

Şimdi {student_name} için {aud['title']} RAPORUNU yukarıdaki sıralama ve kurallara uygun olarak oluştur."""

    return prompt


def app():
    # --- CSS ---
    st.markdown("""
    <style>
        /* ===== GLOBAL DARK MODE KORUMA (öğretmen paneli) ===== */
        .stApp p, .stApp span, .stApp label,
        .stApp h1, .stApp h2, .stApp h3, .stApp h4, .stApp li,
        .stApp [data-testid="stMarkdownContainer"] p,
        .stApp [data-testid="stWidgetLabel"] p {
            color: #0F172A !important;
        }
        [data-testid="stSidebar"] p, [data-testid="stSidebar"] span,
        [data-testid="stSidebar"] label,
        [data-testid="stSidebar"] h1, [data-testid="stSidebar"] h2, [data-testid="stSidebar"] h3 {
            color: #FFFFFF !important;
        }
        .stApp button[kind="primary"] { color: white !important; }
        .stApp .stAlert p { color: inherit !important; }
        .report-header { color: #155724 !important; }
        .id-card-name { color: #1B2A4A !important; }

        /* ===== ÖĞRETMEN PANEL CSS ===== */
        .stSelectbox div, .stMultiSelect div { cursor: pointer !important; }
        div[data-baseweb="select"] { cursor: pointer !important; }
        div[role="listbox"] li { cursor: pointer !important; }
        
        .stRadio > label { 
            font-weight: bold; font-size: 16px; 
            color: #1B2A4A; cursor: pointer !important; 
        }
        .stRadio div[role="radiogroup"] > label { cursor: pointer !important; }
        
        .archive-box { 
            background-color: #f8f9fa; border: 1px solid #ddd; 
            padding: 15px; border-radius: 12px; margin-bottom: 20px; 
        }
        .report-header { 
            color: #155724; background: linear-gradient(135deg, #d4edda, #c3e6cb); 
            padding: 12px 16px; border-radius: 8px; margin-bottom: 10px; 
            border: 1px solid #c3e6cb; font-weight: bold; 
        }
        
        /* Kimlik Kartı */
        .id-card {
            background: #ffffff;
            border: 1px solid #E0E4EA;
            border-radius: 16px;
            padding: 25px;
            border-top: 4px solid #1B2A4A;
            box-shadow: 0 2px 10px rgba(0,0,0,0.05);
        }
        .id-card-name {
            font-size: 1.5rem;
            font-weight: 800;
            color: #1B2A4A;
            margin-bottom: 15px;
        }
    </style>
    """, unsafe_allow_html=True)

    st.markdown("## 👨‍🏫 Öğretmen Yönetim Paneli")
    st.caption("EĞİTİM CHECK UP — Kişisel Eğitim & Kariyer Analiz Merkezi")

    # Giriş ekranına dön butonu
    if st.button("🚪 Giriş Ekranına Dön", key="teacher_to_login"):
        st.session_state.clear()
        st.rerun()

    # Kalıcı veritabanı uyarısı
    if is_using_sqlite():
        st.warning(
            "⚠️ **Veri Kalıcılığı Uyarısı:** Şu an SQLite (geçici) veritabanı kullanılıyor. "
            "Uygulama yeniden başladığında tüm öğrenci verileri silinecektir. "
            "Kalıcı veri için Streamlit Secrets'a `SUPABASE_DB_URL` ekleyin."
        )

    st.markdown("---")

    # Veritabanından verileri çek
    data = get_all_students_with_results()
    student_names_all = [d["info"].name for d in data] if data else []

    # --- SIDEBAR: YÖNETİM ---
    with st.sidebar:
        st.markdown("### ⚙️ Yönetim Araçları")

        # ── VERİ DIŞA AKTARMA ──
        with st.expander("📊 Veri Dışa Aktar (Excel)", expanded=False):
            st.info("Tüm öğrenci bilgilerini, test sonuçlarını ve AI raporlarını tek bir Excel dosyası olarak indirin.")
            if data:
                excel_buffer = generate_full_system_excel(data)
                timestamp = datetime.now().strftime("%Y%m%d_%H%M")
                st.download_button(
                    label="📥 TÜM SİSTEM VERİSİNİ İNDİR",
                    data=excel_buffer,
                    file_name=f"Egitim_CheckUp_TumVeri_{timestamp}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    type="primary",
                    key="dl_full_export"
                )
                st.caption(f"📊 {len(data)} öğrenci · 4 sayfalı Excel dosyası")
            else:
                st.warning("Sistemde henüz veri yok.")

        with st.expander("📄 Tüm PDF Raporları İndir", expanded=False):
            st.info("Her öğrencinin bireysel PDF raporunu tek bir ZIP dosyası olarak indirin.")
            if data:
                if st.button("📄 PDF'leri Oluştur", type="primary", key="btn_bulk_pdf"):
                    import zipfile
                    import io as _io
                    from pdf_engine import generate_student_pdf, generate_student_pdf_filename

                    zip_buffer = _io.BytesIO()
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    success_count = 0

                    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
                        for i, student in enumerate(data):
                            name = student["info"].name
                            status_text.text(f"📄 {name} raporu oluşturuluyor...")
                            try:
                                history = get_student_analysis_history(student["info"].id)
                                pdf_buf = generate_student_pdf(student, history)
                                fname = generate_student_pdf_filename(name)
                                zf.writestr(fname, pdf_buf.getvalue())
                                success_count += 1
                            except Exception as e:
                                st.warning(f"⚠️ {name}: PDF oluşturulamadı — {e}")
                            progress_bar.progress((i + 1) / len(data))

                    zip_buffer.seek(0)
                    status_text.text(f"✅ {success_count}/{len(data)} rapor hazır!")
                    progress_bar.empty()

                    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
                    st.download_button(
                        label=f"📥 {success_count} PDF İNDİR (ZIP)",
                        data=zip_buffer,
                        file_name=f"Egitim_CheckUp_PDF_Raporlar_{timestamp}.zip",
                        mime="application/zip",
                        key="dl_bulk_pdf"
                    )
            else:
                st.warning("Sistemde henüz veri yok.")

        with st.expander("📝 Tüm Word Raporları İndir", expanded=False):
            st.info("Her öğrencinin bireysel Word raporunu tek bir ZIP dosyası olarak indirin.")
            if data:
                if st.button("📝 Word Dosyalarını Oluştur", type="primary", key="btn_bulk_docx"):
                    import zipfile
                    import io as _io
                    from docx_engine import generate_student_docx, generate_student_docx_filename

                    zip_buffer = _io.BytesIO()
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    success_count = 0

                    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
                        for i, student in enumerate(data):
                            name = student["info"].name
                            status_text.text(f"📝 {name} Word raporu oluşturuluyor...")
                            try:
                                history = get_student_analysis_history(student["info"].id)
                                docx_buf = generate_student_docx(student, history)
                                fname = generate_student_docx_filename(name)
                                zf.writestr(fname, docx_buf.getvalue())
                                success_count += 1
                            except Exception as e:
                                st.warning(f"⚠️ {name}: Word oluşturulamadı — {e}")
                            progress_bar.progress((i + 1) / len(data))

                    zip_buffer.seek(0)
                    status_text.text(f"✅ {success_count}/{len(data)} Word raporu hazır!")
                    progress_bar.empty()

                    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
                    st.download_button(
                        label=f"📥 {success_count} WORD İNDİR (ZIP)",
                        data=zip_buffer,
                        file_name=f"Egitim_CheckUp_Word_Raporlar_{timestamp}.zip",
                        mime="application/zip",
                        key="dl_bulk_docx"
                    )
            else:
                st.warning("Sistemde henüz veri yok.")

        with st.expander("🗑️ Öğrenci Dosyası Sil"):
            if not student_names_all:
                st.info("Sistemde kayıtlı öğrenci yok.")
            else:
                st.warning("Seçilen öğrencilerin tüm verileri silinecektir.")
                selected_to_delete = st.multiselect("Silinecekleri Seç:", options=student_names_all)
                if selected_to_delete:
                    if st.button("SEÇİLENLERİ SİL", type="primary"):
                        if delete_specific_students(selected_to_delete):
                            st.success("Kayıtlar silindi.")
                            time.sleep(1)
                            st.rerun()
                        else:
                            st.error("Silme başarısız.")

        st.markdown("---")

        with st.expander("⚠️ Fabrika Ayarlarına Dön"):
            st.error("DİKKAT: Tüm veritabanı silinir!")
            confirm_reset = st.checkbox("Evet, tüm verilerin silineceğini anlıyorum")
            if confirm_reset:
                if st.button("TÜM SİSTEMİ SIFIRLA", type="primary"):
                    if reset_database():
                        st.success("Sistem sıfırlandı.")
                        time.sleep(1)
                        st.rerun()
            else:
                st.info("Devam etmek için onay kutucuğunu işaretleyin.")

    # --- ANA EKRAN ---
    if not data:
        st.info("📂 Henüz kayıtlı öğrenci verisi bulunmamaktadır.")
        return

    # --- GENEL İSTATİSTİKLER ---
    total_students = len(data)
    total_tests = sum(len(d["tests"]) for d in data)
    
    mc1, mc2, mc3 = st.columns(3)
    mc1.metric("👥 Toplam Öğrenci", total_students)
    mc2.metric("📝 Toplam Test", total_tests)
    mc3.metric("📊 Ort. Test/Öğrenci", round(total_tests / total_students, 1) if total_students > 0 else 0)
    
    st.markdown("---")

    # 1. ÖĞRENCİ SEÇİMİ
    st.subheader("📂 Öğrenci Dosyası Görüntüle")

    col1, col2 = st.columns([1, 2])
    with col1:
        selected_name = st.selectbox(
            "Öğrenci Seçiniz:",
            student_names_all,
            index=None,
            placeholder="Listeden bir öğrenci seçin..."
        )

    if not selected_name:
        st.info("👆 Lütfen analizlerini görmek istediğiniz öğrenciyi seçiniz.")
        return

    # Seçilen öğrenci verilerini bul
    student_data = next(d for d in data if d["info"].name == selected_name)
    info = student_data["info"]
    tests = student_data["tests"]

    # 2. ÖĞRENCİ KLASÖRÜ
    st.markdown(f"""
        <div class="id-card">
            <div class="id-card-name">📁 {info.name} — Öğrenci Klasörü</div>
        </div>
    """, unsafe_allow_html=True)

    # Entegre rapor bilgilendirme
    if tests and len(tests) >= 2:
        st.markdown("""
        <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                    padding: 10px 16px; border-radius: 8px; margin-bottom: 10px;">
            <span style="color: white !important; font-size: 14px;">
                🆕 <b>Yeni:</b> "📊 Entegre Rapor (3'lü)" sekmesinden Öğretmen, Öğrenci ve Ebeveyn için
                ayrı ayrı sıralı profesyonel raporlar oluşturabilirsiniz.
            </span>
        </div>
        """, unsafe_allow_html=True)

    tab_profil, tab_testler, tab_ai, tab_entegre = st.tabs([
        "📋 Kişisel Bilgiler",
        "📝 Test Sonuçları",
        "🤖 AI Analiz Raporları",
        "📊 Entegre Rapor (3'lü)"
    ])

    # Öğrenci dosyası indirme butonları (tab'ların üstünde)
    with st.container():
        col_dl1, col_dl2, col_dl3, col_dl4 = st.columns([2, 1, 1, 1])
        with col_dl2:
            student_history = get_student_analysis_history(info.id)
            # PDF İndirme Butonu
            from pdf_engine import generate_student_pdf, generate_student_pdf_filename
            pdf_buffer = generate_student_pdf(student_data, student_history)
            pdf_filename = generate_student_pdf_filename(info.name)
            st.download_button(
                label="📄 PDF Rapor İndir",
                data=pdf_buffer,
                file_name=pdf_filename,
                mime="application/pdf",
                key="dl_student_pdf",
                type="primary"
            )
        with col_dl3:
            # Word İndirme Butonu
            from docx_engine import generate_student_docx, generate_student_docx_filename
            docx_buffer = generate_student_docx(student_data, student_history)
            docx_filename = generate_student_docx_filename(info.name)
            st.download_button(
                label="📝 Word Rapor İndir",
                data=docx_buffer,
                file_name=docx_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_student_docx",
                type="primary"
            )
        with col_dl4:
            student_excel = generate_student_excel(student_data, student_history)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M")
            safe_name = info.name.replace(" ", "_")
            st.download_button(
                label="📥 Excel Dosya İndir",
                data=student_excel,
                file_name=f"{safe_name}_Dosya_{timestamp}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="dl_student_export"
            )

    # ============================================================
    # TAB 1: KİŞİSEL BİLGİLER
    # ============================================================
    with tab_profil:
        st.markdown("#### 👤 Öğrenci Profili")

        grade_val = getattr(info, 'grade', None)
        grade_text = f"{grade_val}. Sınıf" if grade_val else "Belirtilmemiş"

        col_left, col_right = st.columns(2)
        with col_left:
            st.markdown(f"""
            | | |
            |---|---|
            | **👤 Ad Soyad** | {info.name} |
            | **🎂 Yaş** | {info.age} |
            | **⚧ Cinsiyet** | {info.gender} |
            | **🎓 Sınıf** | {grade_text} |
            | **📧 E-posta** | {info.username} |
            """)

        with col_right:
            st.metric("🔑 Toplam Giriş", info.login_count)
            st.metric("📊 Çözülen Test", len(tests))
            if tests:
                son_test = tests[0]
                st.metric("📅 Son Test", f"{son_test['test_name']}")
                st.caption(f"Tarih: {son_test['date']}")

    # ============================================================
    # TAB 2: TAMAMLANAN TESTLER VE OTOMATİK SONUÇLAR
    # ============================================================
    with tab_testler:
        st.markdown("#### 📝 Tamamlanan Testler ve Otomatik Sonuçlar")
        st.caption("Öğrencinin bitirdiği testlerin anlık sistem raporlarını (yapay zekasız) görebilirsiniz.")

        if not tests:
            st.warning("⚠️ Bu öğrenci henüz hiç test çözmemiş.")
        else:
            for idx, t in enumerate(tests):
                btn_label = f"✅ {t['test_name']} (Tarih: {t['date']})"
                with st.expander(btn_label):
                    if t['scores']:
                        fig = plot_scores(t['scores'], t['test_name'])
                        if fig:
                            st.pyplot(fig)

                    st.markdown("### 📄 Sistem Raporu")
                    if t.get('report'):
                        st.markdown(t['report'])
                    else:
                        st.warning("Bu test için otomatik rapor bulunamadı.")
                        st.write("Ham Cevaplar:", t['raw_answers'])

    # ============================================================
    # TAB 3: KAYITLI AI RAPOR ARŞİVİ
    # ============================================================
    with tab_ai:
        st.markdown("#### 📂 Kayıtlı AI Rapor Arşivi")
        st.caption("Daha önce Claude ile oluşturduğunuz detaylı analizler.")

        history = student_history  # Zaten yukarıda çekildi

        if not history:
            st.info("Bu öğrenci için henüz AI destekli analiz raporu oluşturulmamış.")
        else:
            st.markdown(f"**{len(history)} adet** kayıtlı rapor bulundu.")

            for idx, record in enumerate(history):
                btn_label = f"🤖 AI Raporu {idx+1}: {record['combination']} ({record['date']})"
                with st.expander(btn_label):
                    st.markdown(f"<div class='report-header'>ANALİZ KAPSAMI: {record['combination']}</div>", unsafe_allow_html=True)

                    archived_test_names = record['combination'].split(' + ')
                    archived_test_data = [t for t in tests if t["test_name"] in archived_test_names]

                    if archived_test_data:
                        st.markdown("#### 📊 Grafik Özeti")
                        g_cols = st.columns(2)
                        for i, t_data in enumerate(archived_test_data):
                            if t_data["scores"]:
                                fig = plot_scores(t_data["scores"], t_data["test_name"])
                                if fig:
                                    g_cols[i % 2].pyplot(fig)
                        st.markdown("---")

                    st.markdown(record['report'])
                    st.download_button(
                        label=f"📥 Raporu İndir ({idx+1})",
                        data=record['report'],
                        file_name=f"{info.name}_AI_Rapor_{idx+1}.txt",
                        mime="text/plain",
                        key=f"dl_{idx}"
                    )

        st.divider()

        # ============================================================
        # 4.5 VELİ ÖZETİ OLUŞTURMA
        # ============================================================
        st.subheader("👨‍👩‍👦 Veli Özeti Oluştur")
        st.caption("Kayıtlı AI raporlarından seçtiğiniz başlıkları veliye sunmak üzere kısa ve anlaşılır bir özet haline getirin.")

        if not history:
            st.info("Veli özeti oluşturabilmek için önce en az bir AI analiz raporu oluşturmanız gerekir.")
        else:
            # Rapor seçimi
            report_options = [f"Rapor {idx+1}: {rec['combination']} ({rec['date']})" for idx, rec in enumerate(history)]
            selected_report_label = st.selectbox(
                "📄 Hangi rapordan özet oluşturmak istiyorsunuz?",
                options=report_options,
                key="veli_ozet_rapor_sec"
            )

            if selected_report_label:
                selected_report_idx = report_options.index(selected_report_label)
                selected_report_content = history[selected_report_idx]['report']

                # Rapordaki başlıkları otomatik çıkar
                import re
                headings = re.findall(r'^#{1,3}\s+(.+)$', selected_report_content, re.MULTILINE)
                if not headings:
                    headings = re.findall(r'^\*\*(.+?)\*\*', selected_report_content, re.MULTILINE)

                if headings:
                    st.write("📌 **Veliye sunmak istediğiniz başlıkları seçiniz:**")
                    selected_headings = st.multiselect(
                        "Başlık Listesi:",
                        options=headings,
                        default=None,
                        key="veli_ozet_basliklar",
                        help="Birden fazla başlık seçebilirsiniz. Seçtiğiniz başlıklar altındaki içerik, veli için yalın ve kısa bir dilde özetlenecektir."
                    )

                    if selected_headings:
                        # Ek notlar
                        teacher_notes = st.text_area(
                            "📝 Öğretmen Notu (İsteğe bağlı):",
                            placeholder="Veliye iletmek istediğiniz ek notlarınız varsa buraya yazın...",
                            key="veli_ozet_not",
                            height=80
                        )

                        if st.button("📋 Veli Özeti Oluştur", type="primary", key="btn_veli_ozet"):
                            # Seçilen başlıkların altındaki içerikleri çıkar
                            sections_content = []
                            for heading in selected_headings:
                                escaped = re.escape(heading)
                                pattern = rf'#{{{1,3}}}\s+{escaped}\s*\n(.*?)(?=\n#{{{1,3}}}\s|\Z)'
                                match = re.search(pattern, selected_report_content, re.DOTALL)
                                if match:
                                    sections_content.append(f"### {heading}\n{match.group(1).strip()}")
                                else:
                                    # Fallback: başlığı ve sonraki birkaç paragrafı al
                                    idx_pos = selected_report_content.find(heading)
                                    if idx_pos >= 0:
                                        snippet = selected_report_content[idx_pos:idx_pos+2000]
                                        sections_content.append(f"### {heading}\n{snippet}")

                            combined_sections = "\n\n".join(sections_content)
                            teacher_note_section = f"\n\nÖğretmenin ek notu: {teacher_notes}" if teacher_notes.strip() else ""

                            veli_prompt = f"""Aşağıda bir öğrencinin psikometrik test analiz raporundan seçilmiş bölümler bulunmaktadır. Bu bölümleri, öğrencinin velisine sunulmak üzere kısa ve anlaşılır bir özet haline getir.

# KRİTİK KURALLAR:
1. Veli, eğitim psikolojisi terminolojisine hakim değildir. Her ifadeyi günlük konuşma Türkçesiyle yaz.
2. Bilimsel terimler kullanma. Zorunlu olursa parantez içinde çok kısa ve basit açıklama ekle.
3. Abartılı övgü veya motivasyon ifadelerinden kaçın. Gerçekçi, samimi ve yapıcı ol.
4. Her başlık için en fazla 3-4 cümlelik bir özet yaz. Tüm özet toplamda 1 sayfa A4'ü geçmesin.
5. Özeti "Sayın Veli," diye başlat ve profesyonel ama sıcak bir dille yaz.
6. Sonunda velinin evde yapabileceği en önemli 3 somut adımı kısa maddeler halinde listele.
7. Öğrencinin adı: {info.name}

# SEÇİLMİŞ RAPOR BÖLÜMLERİ:
{combined_sections}
{teacher_note_section}

Şimdi bu bilgileri veliye uygun kısa bir özete dönüştür."""

                            with st.spinner("Veli özeti hazırlanıyor..."):
                                veli_summary = get_ai_analysis(veli_prompt)

                            st.success("✅ Veli özeti hazır!")
                            st.markdown("---")
                            st.markdown("#### 📋 Veli Özeti")
                            st.markdown(veli_summary)

                            # İndirme butonları
                            col_vdl1, col_vdl2 = st.columns(2)
                            with col_vdl1:
                                st.download_button(
                                    label="📥 Veli Özetini İndir (TXT)",
                                    data=veli_summary,
                                    file_name=f"{info.name}_Veli_Ozeti.txt",
                                    mime="text/plain",
                                    key="dl_veli_ozet_txt"
                                )
                            with col_vdl2:
                                # Word formatında da sunalım
                                try:
                                    from docx_engine import generate_student_docx
                                    # Basit bir Word belgesi oluştur
                                    from docx import Document
                                    from io import BytesIO
                                    doc = Document()
                                    doc.add_heading(f"{info.name} — Veli Bilgilendirme Özeti", level=1)
                                    doc.add_paragraph(f"Tarih: {datetime.now().strftime('%d.%m.%Y')}")
                                    doc.add_paragraph("")
                                    for para in veli_summary.split('\n'):
                                        if para.strip():
                                            clean_para = para.replace('**', '').replace('###', '').replace('##', '').replace('#', '').strip()
                                            if clean_para:
                                                doc.add_paragraph(clean_para)
                                    buf = BytesIO()
                                    doc.save(buf)
                                    buf.seek(0)
                                    st.download_button(
                                        label="📝 Veli Özetini İndir (Word)",
                                        data=buf,
                                        file_name=f"{info.name}_Veli_Ozeti.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        key="dl_veli_ozet_docx"
                                    )
                                except ImportError:
                                    pass
                else:
                    st.warning("Seçilen raporda otomatik olarak başlık bulunamadı. Lütfen farklı bir rapor deneyin.")

        st.divider()

        # ============================================================
        # 5. YENİ AI ANALİZİ OLUŞTURMA
        # ============================================================
        st.subheader("⚡ Yeni AI Analizi Oluştur")

        if not tests:
            st.write("Analiz yapılacak veri yok.")
        else:
            all_completed_tests = [t["test_name"] for t in tests]

            st.write("Analiz raporu oluşturmak istediğiniz testleri seçiniz:")
            selected_tests = st.multiselect(
                "Test Listesi:",
                options=all_completed_tests,
                default=all_completed_tests
            )

            if selected_tests:
                st.markdown("---")
                st.write("📊 **Analiz Yöntemini Seçiniz:**")

                analysis_mode = st.radio(
                    "Nasıl bir rapor istiyorsunuz?",
                    options=["BÜTÜNCÜL (Harmanlanmış) Rapor", "AYRI AYRI (Tekil) Raporlar"],
                    index=0,
                    help="Bütüncül: Seçilen tüm testleri birleştirip 'Büyük Resim' sentezi yapar.\nAyrı Ayrı: Seçilen her test için sırayla detaylı psikometrik analiz yapar."
                )

                st.markdown("<br>", unsafe_allow_html=True)

                # ── RAPOR TEKRAR OLUŞTURMA KORUMASI ──
                existing_combinations = [rec['combination'] for rec in history] if history else []
                
                if analysis_mode == "BÜTÜNCÜL (Harmanlanmış) Rapor":
                    check_comb = " + ".join(sorted(selected_tests))
                    already_exists = any(
                        set(ec.split(" + ")) == set(selected_tests) for ec in existing_combinations
                    )
                else:
                    already_exists = any(
                        t_name in existing_combinations for t_name in selected_tests
                    )
                    existing_singles = [t_name for t_name in selected_tests if t_name in existing_combinations]

                confirm_key = f"confirm_regen_{info.id}_{analysis_mode}"
                
                if already_exists:
                    if analysis_mode == "BÜTÜNCÜL (Harmanlanmış) Rapor":
                        st.warning(
                            "⚠️ **Bu test kombinasyonu için daha önce bir bütüncül rapor oluşturulmuş.**\n\n"
                            "Tekrar oluşturmak, arşive yeni bir kayıt ekleyecektir. Devam etmek istediğinizden emin misiniz?"
                        )
                    else:
                        st.warning(
                            f"⚠️ **Şu testler için daha önce tekil rapor oluşturulmuş:** "
                            f"{', '.join(existing_singles)}\n\n"
                            "Tekrar oluşturmak, arşive yeni kayıtlar ekleyecektir. Devam etmek istediğinizden emin misiniz?"
                        )
                    
                    col_confirm, col_cancel = st.columns(2)
                    with col_confirm:
                        proceed = st.button("✅ Evet, Eminim — Yeniden Oluştur", type="primary", key=confirm_key)
                    with col_cancel:
                        cancel = st.button("❌ Hayır, Vazgeç", key=f"cancel_{confirm_key}")
                    
                    if cancel:
                        st.info("İşlem iptal edildi.")
                        proceed = False
                else:
                    proceed = st.button("🚀 ANALİZİ BAŞLAT (Claude AI)", type="primary")

                if proceed:
                    analyzed_data = [t for t in tests if t["test_name"] in selected_tests]

                    # Grafikleri göster
                    st.markdown("### 📊 Puan Grafikleri")
                    gc = st.columns(2)
                    for i, t in enumerate(analyzed_data):
                        if t["scores"]:
                            fig = plot_scores(t["scores"], t["test_name"])
                            if fig:
                                gc[i % 2].pyplot(fig)
                            else:
                                gc[i % 2].info(f"{t['test_name']} için grafik verisi yok.")

                    # ====================================================
                    # MOD 1: BÜTÜNCÜL ANALİZ
                    # ====================================================
                    if analysis_mode == "BÜTÜNCÜL (Harmanlanmış) Rapor":
                        st.info(f"⏳ Claude AI, seçilen **{len(selected_tests)} testi** harmanlıyor...")

                        with st.spinner("Stratejik analiz hazırlanıyor..."):
                            ai_input = []
                            for t in analyzed_data:
                                raw = t.get("raw_answers", "")
                                if isinstance(raw, str):
                                    try:
                                        raw = json.loads(raw)
                                    except (json.JSONDecodeError, ValueError):
                                        raw = raw

                                ai_input.append({
                                    "TEST_ADI": t["test_name"],
                                    "TARİH": str(t["date"]),
                                    "SONUÇLAR": t["scores"] if t["scores"] else raw
                                })

                            prompt = build_holistic_prompt(
                                student_name=info.name,
                                student_age=info.age,
                                student_gender=info.gender,
                                test_data_list=ai_input,
                                student_grade=getattr(info, 'grade', None)
                            )

                            final_report = get_ai_analysis(prompt)
                            save_holistic_analysis(info.id, selected_tests, final_report)

                            st.success("✅ Bütüncül analiz tamamlandı ve arşive kaydedildi.")
                            time.sleep(1.5)
                            st.rerun()

                    # ====================================================
                    # MOD 2: AYRI AYRI TEKİL ANALİZLER
                    # ====================================================
                    else:
                        progress_text = "Testler sırayla analiz ediliyor..."
                        my_bar = st.progress(0, text=progress_text)
                        total_ops = len(analyzed_data)

                        for idx, t in enumerate(analyzed_data):
                            test_name = t["test_name"]
                            my_bar.progress(
                                (idx + 1) / total_ops,
                                text=f"**{test_name}** analiz ediliyor... ({idx+1}/{total_ops})"
                            )

                            raw = t.get("raw_answers", "")
                            if isinstance(raw, str):
                                try:
                                    raw = json.loads(raw)
                                except (json.JSONDecodeError, ValueError):
                                    raw = raw

                            test_data_for_prompt = {
                                "TEST_ADI": test_name,
                                "TARİH": str(t["date"]),
                                "SONUÇLAR": t["scores"] if t["scores"] else raw
                            }

                            prompt = build_single_test_prompt(
                                student_name=info.name,
                                student_age=info.age,
                                student_gender=info.gender,
                                test_name=test_name,
                                test_data=test_data_for_prompt,
                                student_grade=getattr(info, 'grade', None)
                            )

                            single_report = get_ai_analysis(prompt)
                            save_holistic_analysis(info.id, [test_name], single_report)

                        my_bar.empty()
                        st.success(f"✅ {total_ops} test başarıyla analiz edildi ve Arşiv'e eklendi.")
                        time.sleep(2)
                        st.rerun()

    # ============================================================
    # TAB 4: ENTEGRE RAPOR (3'LÜ) — Öğretmen / Öğrenci / Ebeveyn
    # ============================================================
    with tab_entegre:
        st.markdown("""
        <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                    padding: 20px; border-radius: 12px; margin-bottom: 20px;">
            <h3 style="color: white !important; margin: 0;">📊 Entegre Rapor Sistemi</h3>
            <p style="color: rgba(255,255,255,0.9) !important; margin: 5px 0 0 0;">
                Tüm testleri belirli bir sıra ile analiz edip <b>Öğretmen</b>, <b>Öğrenci</b> ve <b>Ebeveyn</b>
                için ayrı ayrı profesyonel raporlar oluşturur.
            </p>
        </div>
        """, unsafe_allow_html=True)

        if not tests:
            st.warning("⚠️ Bu öğrenci henüz hiç test çözmemiş. Entegre rapor oluşturulamaz.")
        else:
            # Testleri patronun istediği sıraya göre sırala
            ordered_tests = _order_tests(tests)
            completed_names = [t["test_name"] for t in ordered_tests]
            canonical_completed = [_normalize_test_name(n) for n in completed_names]

            # Hangi testler mevcut, hangileri eksik göster
            st.markdown("#### 📋 Test Durumu (Sıralı)")
            col_pot, col_mev = st.columns(2)

            with col_pot:
                st.markdown("**🔬 Potansiyel Analiz Testleri:**")
                for t_name in INTEGRATED_TEST_ORDER[:5]:
                    if t_name in canonical_completed:
                        st.markdown(f"✅ {t_name}")
                    else:
                        st.markdown(f"⬜ {t_name} *(çözülmemiş)*")

            with col_mev:
                st.markdown("**📈 Mevcut Durum Testleri:**")
                for t_name in INTEGRATED_TEST_ORDER[5:]:
                    if t_name in canonical_completed:
                        st.markdown(f"✅ {t_name}")
                    else:
                        st.markdown(f"⬜ {t_name} *(çözülmemiş)*")

            available_count = sum(1 for t in INTEGRATED_TEST_ORDER if t in canonical_completed)
            st.info(f"📊 **{available_count}/{len(INTEGRATED_TEST_ORDER)}** test çözülmüş. Çözülmemiş testler raporda atlanacaktır.")

            st.markdown("---")

            # Rapor türü seçimi
            st.markdown("#### 🎯 Rapor Oluştur")
            report_type_selection = st.radio(
                "Hangi raporları oluşturmak istiyorsunuz?",
                options=[
                    "📚 Üçünü Birden (Öğretmen + Öğrenci + Ebeveyn)",
                    "👨‍🏫 Sadece Öğretmen/Koç Raporu",
                    "🎓 Sadece Öğrenci Raporu",
                    "👨‍👩‍👦 Sadece Ebeveyn Raporu"
                ],
                index=0,
                key="entegre_rapor_tipi"
            )

            # Mevcut entegre raporları kontrol et
            entegre_history = get_student_analysis_history(info.id)
            existing_entegre = [
                r for r in entegre_history
                if r['combination'].startswith("ENTEGRE_")
            ] if entegre_history else []

            if existing_entegre:
                st.markdown("---")
                st.markdown("#### 📂 Mevcut Entegre Raporlar")
                for idx, rec in enumerate(existing_entegre):
                    rapor_label = rec['combination'].replace("ENTEGRE_", "").replace("_", " ").title()
                    with st.expander(f"📄 {rapor_label} Raporu ({rec['date']})"):
                        st.markdown(rec['report'])

                        col_d1, col_d2 = st.columns(2)
                        with col_d1:
                            st.download_button(
                                label="📥 İndir (TXT)",
                                data=rec['report'],
                                file_name=f"{info.name}_{rapor_label}_Entegre_Rapor.txt",
                                mime="text/plain",
                                key=f"dl_entegre_{idx}"
                            )
                        with col_d2:
                            try:
                                from docx import Document
                                from io import BytesIO
                                doc = Document()
                                doc.add_heading(f"{info.name} — {rapor_label} Entegre Raporu", level=1)
                                doc.add_paragraph(f"Tarih: {rec['date']}")
                                doc.add_paragraph("")
                                for para in rec['report'].split('\n'):
                                    if para.strip():
                                        clean = para.replace('**', '').replace('###', '').replace('##', '').replace('#', '').strip()
                                        if clean:
                                            doc.add_paragraph(clean)
                                buf = BytesIO()
                                doc.save(buf)
                                buf.seek(0)
                                st.download_button(
                                    label="📝 İndir (Word)",
                                    data=buf,
                                    file_name=f"{info.name}_{rapor_label}_Entegre_Rapor.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key=f"dl_entegre_docx_{idx}"
                                )
                            except ImportError:
                                pass

                st.markdown("---")

            # Oluştur butonu
            if available_count < 2:
                st.warning("En az 2 test çözülmüş olmalıdır. Lütfen öğrencinin daha fazla test çözmesini bekleyin.")
            else:
                if st.button("🚀 ENTEGRE RAPOR OLUŞTUR", type="primary", key="btn_entegre_rapor"):
                    # Test verilerini hazırla (sıralı)
                    ai_input = []
                    for t in ordered_tests:
                        raw = t.get("raw_answers", "")
                        if isinstance(raw, str):
                            try:
                                raw = json.loads(raw)
                            except (json.JSONDecodeError, ValueError):
                                raw = raw
                        ai_input.append({
                            "TEST_ADI": t["test_name"],
                            "TARİH": str(t["date"]),
                            "SONUÇLAR": t["scores"] if t["scores"] else raw
                        })

                    # Hangi raporlar oluşturulacak
                    type_map = {
                        "📚 Üçünü Birden (Öğretmen + Öğrenci + Ebeveyn)": [
                            ("ogretmen", "Öğretmen/Koç"),
                            ("ogrenci", "Öğrenci"),
                            ("ebeveyn", "Ebeveyn")
                        ],
                        "👨‍🏫 Sadece Öğretmen/Koç Raporu": [("ogretmen", "Öğretmen/Koç")],
                        "🎓 Sadece Öğrenci Raporu": [("ogrenci", "Öğrenci")],
                        "👨‍👩‍👦 Sadece Ebeveyn Raporu": [("ebeveyn", "Ebeveyn")]
                    }
                    selected_types = type_map[report_type_selection]

                    total_reports = len(selected_types)
                    progress_bar = st.progress(0, text="Raporlar hazırlanıyor...")

                    for i, (rtype, rlabel) in enumerate(selected_types):
                        progress_bar.progress(
                            (i) / total_reports,
                            text=f"📝 **{rlabel}** raporu oluşturuluyor... ({i+1}/{total_reports})"
                        )

                        prompt = build_integrated_report_prompt(
                            student_name=info.name,
                            student_age=info.age,
                            student_gender=info.gender,
                            test_data_list=ai_input,
                            report_type=rtype,
                            student_grade=getattr(info, 'grade', None)
                        )

                        report_text = get_ai_analysis(prompt)
                        combination_key = f"ENTEGRE_{rlabel.upper().replace('/', '_')}"
                        save_holistic_analysis(info.id, [combination_key], report_text)

                    progress_bar.progress(1.0, text="✅ Tüm raporlar tamamlandı!")
                    time.sleep(1)
                    st.success(f"✅ {total_reports} adet entegre rapor oluşturuldu ve arşive kaydedildi!")
                    time.sleep(1.5)
                    st.rerun()

    # 6. HAM VERİ LİSTESİ
    st.divider()
    with st.expander("🗂️ Ham Veri Listesi"):
        if tests:
            import pandas as pd
            df_tests = pd.DataFrame(tests)
            df_tests['date'] = pd.to_datetime(df_tests['date'], errors='coerce').dt.strftime('%d.%m.%Y')
            st.dataframe(df_tests[["test_name", "date"]], use_container_width=True)
