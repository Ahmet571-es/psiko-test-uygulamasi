"""
Word (DOCX) Rapor Motoru — Eğitim Check-Up
Öğrenci kişisel bilgileri + test sonuçları + grafikler + AI analiz raporlarını
profesyonel Word formatında dışa aktarır.

v1.0 — python-docx ile Türkçe uyumlu Word rapor üretimi
"""

import io
import os
import re
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============================================================
# RENK TANIMLARI
# ============================================================
CLR_PRIMARY   = RGBColor(0x1B, 0x2A, 0x4A)  # Koyu lacivert
CLR_SECONDARY = RGBColor(0x2E, 0x86, 0xAB)  # Açık mavi
CLR_ACCENT    = RGBColor(0xA2, 0x3B, 0x72)  # Bordo/mor
CLR_GRAY      = RGBColor(0x6C, 0x75, 0x7D)  # Gri
CLR_DARK      = RGBColor(0x21, 0x25, 0x29)  # Koyu metin
CLR_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)


# ============================================================
# YARDIMCI FONKSİYONLAR
# ============================================================

def _set_cell_shading(cell, color_hex):
    """Hücre arka plan rengini ayarla."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _set_cell_text(cell, text, bold=False, color=None, size=9, alignment=None):
    """Hücre metnini formatlı yaz."""
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color


def _add_styled_heading(doc, text, level=1, color=None):
    """Renkli başlık ekle."""
    heading = doc.add_heading(text, level=level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = color
    return heading


def _add_colored_bar(doc, text, bg_color_hex="1B2A4A"):
    """Renkli arka planlı başlık çubuğu (tablo olarak)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _set_cell_shading(cell, bg_color_hex)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = CLR_WHITE
    run.font.name = "Calibri"
    # Tablo kenarlıklarını kaldır
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tbl_pr.append(borders)
    return table


def _add_markdown_content(doc, md_text):
    """Basit Markdown → Word paragraf dönüştürücü."""
    if not md_text:
        doc.add_paragraph("Rapor içeriği bulunamadı.", style="Body Text").italic = True
        return

    lines = md_text.split('\n')
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Heading'ler
        if stripped.startswith('#### '):
            text = re.sub(r'[*#]', '', stripped[5:]).strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = CLR_ACCENT
            continue
        if stripped.startswith('### '):
            text = re.sub(r'[*#]', '', stripped[4:]).strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = CLR_SECONDARY
            continue
        if stripped.startswith('## '):
            text = re.sub(r'[*#]', '', stripped[3:]).strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = CLR_PRIMARY
            continue
        if stripped.startswith('# '):
            text = re.sub(r'[*#]', '', stripped[2:]).strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = CLR_PRIMARY
            continue

        # Yatay çizgi
        if stripped in ('---', '***', '___'):
            doc.add_paragraph('─' * 60)
            continue

        # Madde işaretleri
        bullet_match = re.match(r'^[-*•]\s+(.+)', stripped)
        if bullet_match:
            text = bullet_match.group(1)
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            p = doc.add_paragraph(style='List Bullet')
            _add_inline_runs(p, text)
            continue

        # Numaralı liste
        num_match = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if num_match:
            text = num_match.group(2)
            text_clean = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            p = doc.add_paragraph(style='List Number')
            _add_inline_runs(p, text_clean)
            continue

        # Normal paragraf
        p = doc.add_paragraph()
        _add_inline_runs(p, stripped)


def _add_inline_runs(paragraph, text):
    """Markdown bold/italic → Word runs."""
    # **bold** ve *italic* ayrıştır
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = "Calibri"
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.size = Pt(9)
            run.font.name = "Calibri"
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(9)
            run.font.name = "Calibri"


# ============================================================
# GRAFİK OLUŞTURMA (pdf_engine ile paylaşımlı mantık)
# ============================================================

def _create_chart_for_docx(test_name, scores):
    """Matplotlib grafiği oluştur, BytesIO olarak döndür."""
    try:
        from pdf_engine import _extract_chart_data, _create_chart_image
        plot_data, chart_title = _extract_chart_data(test_name, scores)
        if not plot_data:
            return None
        img_buffer = _create_chart_image(plot_data, chart_title, chart_width_cm=14)
        if img_buffer:
            img_buffer.seek(0)
            return img_buffer
        return None
    except Exception:
        return None


# ============================================================
# ANA WORD ÜRETİCİ
# ============================================================

def generate_student_docx(student_data, analysis_history, include_charts=True):
    """
    Öğrencinin tüm bilgilerini Word formatında dışa aktarır.

    Args:
        student_data: dict — {"info": StudentInfo, "tests": [...]}
        analysis_history: list — AI analiz geçmişi
        include_charts: bool — Grafikleri dahil et

    Returns:
        io.BytesIO — DOCX dosya buffer'ı
    """
    info = student_data["info"]
    tests = student_data["tests"]
    now = datetime.now()

    doc = Document()

    # Varsayılan font ayarı
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # Sayfa kenar boşlukları
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # ──────────────────────────────────────
    # KAPAK SAYFASI
    # ──────────────────────────────────────
    doc.add_paragraph()  # Boşluk

    # Başlık
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("EĞİTİM CHECK-UP")
    title_run.bold = True
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = CLR_PRIMARY
    title_run.font.name = "Calibri"

    # Alt başlık
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Kişisel Eğitim & Kariyer Analiz Merkezi")
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = CLR_SECONDARY
    sub_run.font.name = "Calibri"

    doc.add_paragraph()  # Boşluk

    # Öğrenci adı
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(str(info.name))
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.color.rgb = CLR_PRIMARY
    name_run.font.name = "Calibri"

    # Rapor türü
    type_p = doc.add_paragraph()
    type_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    type_run = type_p.add_run("Psikometrik Değerlendirme Raporu")
    type_run.font.size = Pt(11)
    type_run.font.color.rgb = CLR_SECONDARY
    type_run.font.name = "Calibri"

    doc.add_paragraph()

    # Tarih
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_p.add_run(f"Rapor Tarihi: {now.strftime('%d.%m.%Y %H:%M')}")
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = CLR_GRAY
    date_run.font.name = "Calibri"

    doc.add_page_break()

    # ──────────────────────────────────────
    # BÖLÜM 1: ÖĞRENCİ PROFİLİ
    # ──────────────────────────────────────
    _add_colored_bar(doc, "BÖLÜM 1: ÖĞRENCİ PROFİLİ")
    doc.add_paragraph()

    grade_val = getattr(info, 'grade', None)
    from db_utils import format_grade
    grade_text = format_grade(grade_val)

    profile_rows = [
        ("Ad Soyad", str(info.name)),
        ("Yaş", str(info.age)),
        ("Cinsiyet", str(info.gender)),
        ("Sınıf", grade_text),
        ("E-posta / Kullanıcı", str(info.username)),
        ("Toplam Giriş", str(info.login_count)),
        ("Çözülen Test Sayısı", str(len(tests))),
    ]

    table = doc.add_table(rows=len(profile_rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, value) in enumerate(profile_rows):
        _set_cell_text(table.cell(i, 0), label, bold=True, size=9)
        _set_cell_shading(table.cell(i, 0), "F8F9FA")
        _set_cell_text(table.cell(i, 1), value, size=9)

    # Tablo stili
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)

    # Tamamlanan testler listesi
    if tests:
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("Tamamlanan Testler:")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = CLR_SECONDARY

        test_tbl = doc.add_table(rows=len(tests) + 1, cols=3)
        test_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header
        headers = ["No", "Test Adı", "Tarih"]
        for j, h in enumerate(headers):
            _set_cell_text(test_tbl.cell(0, j), h, bold=True, color=CLR_WHITE, size=9,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_shading(test_tbl.cell(0, j), "1B2A4A")

        for i, t in enumerate(tests, 1):
            _set_cell_text(test_tbl.cell(i, 0), str(i), size=9,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(test_tbl.cell(i, 1), t["test_name"], size=9)
            _set_cell_text(test_tbl.cell(i, 2), str(t.get("date", "-")), size=9,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)

            # Zebra striping
            if i % 2 == 0:
                for j in range(3):
                    _set_cell_shading(test_tbl.cell(i, j), "F8F9FA")

    doc.add_page_break()

    # ──────────────────────────────────────
    # BÖLÜM 2: TEST SONUÇLARI
    # ──────────────────────────────────────
    _add_colored_bar(doc, "BÖLÜM 2: TEST SONUÇLARI VE PUAN DETAYLARI")
    doc.add_paragraph()

    if not tests:
        p = doc.add_paragraph("Henüz test sonucu bulunmamaktadır.")
        p.runs[0].font.color.rgb = CLR_GRAY
    else:
        for idx, t in enumerate(tests):
            # Test başlığı (mavi çubuk)
            _add_colored_bar(doc,
                f"{idx+1}. {t['test_name']}  |  Tarih: {t.get('date', '-')}",
                bg_color_hex="2E86AB")
            doc.add_paragraph()

            # Puan tablosu
            if t.get("scores") and isinstance(t["scores"], dict):
                scores = t["scores"]
                score_tbl = doc.add_table(rows=len(scores) + 1, cols=2)
                score_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

                _set_cell_text(score_tbl.cell(0, 0), "Ölçek / Boyut", bold=True,
                              color=CLR_WHITE, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                _set_cell_shading(score_tbl.cell(0, 0), "1B2A4A")
                _set_cell_text(score_tbl.cell(0, 1), "Puan / Değer", bold=True,
                              color=CLR_WHITE, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                _set_cell_shading(score_tbl.cell(0, 1), "1B2A4A")

                for i, (k, v) in enumerate(scores.items(), 1):
                    _set_cell_text(score_tbl.cell(i, 0), str(k), size=9)
                    _set_cell_text(score_tbl.cell(i, 1), str(v), size=9,
                                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    if i % 2 == 0:
                        _set_cell_shading(score_tbl.cell(i, 0), "F8F9FA")
                        _set_cell_shading(score_tbl.cell(i, 1), "F8F9FA")

                doc.add_paragraph()

                # GRAFİK EKLE
                if include_charts:
                    try:
                        chart_buf = _create_chart_for_docx(t['test_name'], scores)
                        if chart_buf:
                            doc.add_picture(chart_buf, width=Cm(14))
                            last_p = doc.paragraphs[-1]
                            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            doc.add_paragraph()
                    except Exception:
                        pass

            # Sistem raporu
            report_text = t.get("report", "")
            if report_text:
                p = doc.add_paragraph()
                run = p.add_run("Sistem Raporu:")
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = CLR_ACCENT
                _add_markdown_content(doc, report_text)

            # Ayırıcı
            doc.add_paragraph('─' * 60)

    doc.add_page_break()

    # ──────────────────────────────────────
    # BÖLÜM 3: AI ANALİZ RAPORLARI
    # ──────────────────────────────────────
    _add_colored_bar(doc, "BÖLÜM 3: AI DESTEKLİ ANALİZ RAPORLARI (Claude)")
    doc.add_paragraph()

    if not analysis_history:
        p = doc.add_paragraph(
            "Bu öğrenci için henüz AI destekli analiz raporu oluşturulmamıştır."
        )
        p.runs[0].font.color.rgb = CLR_GRAY
    else:
        p = doc.add_paragraph()
        run = p.add_run(f"Toplam {len(analysis_history)} adet AI analiz raporu bulunmaktadır.")
        run.font.size = Pt(9)
        doc.add_paragraph()

        for idx, record in enumerate(analysis_history):
            combo = record.get('combination', 'Bilinmiyor')
            date = record.get('date', '-')

            _add_colored_bar(doc,
                f"AI Rapor {idx+1}: {combo}  |  {date}",
                bg_color_hex="A23B72")
            doc.add_paragraph()

            report_content = record.get('report', '')
            if report_content:
                _add_markdown_content(doc, report_content)
            else:
                p = doc.add_paragraph("Rapor içeriği boş.")
                p.runs[0].italic = True

            doc.add_paragraph('─' * 60)

            if idx < len(analysis_history) - 1:
                doc.add_page_break()

    # ──────────────────────────────────────
    # FOOTER / DİSCLAİMER
    # ──────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph('━' * 60)

    disclaimer = (
        "Bu rapor, EĞİTİM CHECK-UP psikometrik değerlendirme sistemi tarafından, "
        "yapay zekâ destekli analiz altyapısıyla üretilmiştir. Raporda yer alan "
        "tüm yorumlar, öğrencinin psikometrik test verilerine dayanmaktadır. "
        "Bu rapor klinik tanı içermez ve klinik değerlendirme yerine geçmez."
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(disclaimer)
    run.font.size = Pt(7.5)
    run.font.color.rgb = CLR_GRAY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(
        f"Rapor üretim tarihi: {now.strftime('%d.%m.%Y %H:%M')} | Eğitim Check-Up v2.1"
    )
    run2.font.size = Pt(7.5)
    run2.font.color.rgb = CLR_GRAY

    # Footer ekle (sayfa numarası)
    _add_footer(doc)

    # Buffer'a kaydet
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _add_footer(doc):
    """Sayfa numarası footer'ı ekle."""
    try:
        section = doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run1 = p.add_run("Eğitim Check-Up | Sayfa ")
        run1.font.size = Pt(7)
        run1.font.color.rgb = CLR_GRAY

        # Sayfa numarası field
        fld_xml = (
            '<w:fldSimple {} w:instr=" PAGE \\* MERGEFORMAT ">'
            '<w:r><w:rPr><w:sz w:val="14"/></w:rPr><w:t>1</w:t></w:r>'
            '</w:fldSimple>'
        ).format(nsdecls('w'))
        fld_simple = parse_xml(fld_xml)
        p._p.append(fld_simple)
    except Exception:
        pass


def generate_student_docx_filename(student_name):
    """DOCX dosya adı üretir."""
    safe_name = student_name.replace(" ", "_")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    return f"{safe_name}_Rapor_{timestamp}.docx"
