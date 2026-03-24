"""
Aile Bilgilendirme Ozeti — PDF ve DOCX Rapor Uretimi
=====================================================
Mevcut pdf_engine.py ve docx_engine.py altyapisindan turetilmistir.
DejaVu Sans font ile Turkce karakter destegi.
"""

import io
import os
import re
from datetime import datetime

# ============================================================
# PDF URETIMI (ReportLab)
# ============================================================

def generate_family_summary_pdf(student_name, student_age, student_gender,
                                teacher_name, summary_text, test_types,
                                student_grade=None):
    """Aile bilgilendirme ozeti icin PDF olusturur."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor, white, black
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        HRFlowable
    )

    # Font kaydi — pdf_engine ile ayni mantik
    from pdf_engine import FONT_NAME, FONT_BOLD, FONT_ITALIC, _safe

    PRIMARY = HexColor("#1B2A4A")
    SECONDARY = HexColor("#2E86AB")
    ACCENT = HexColor("#A23B72")
    BORDER = HexColor("#DEE2E6")

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=2 * cm, rightMargin=2 * cm,
        topMargin=2 * cm, bottomMargin=2.5 * cm,
        title="Veli Bilgilendirme Ozeti",
    )

    # Stiller
    style_title = ParagraphStyle(
        'FamTitle', fontName=FONT_BOLD, fontSize=18,
        textColor=PRIMARY, alignment=TA_CENTER, spaceAfter=4, leading=24,
    )
    style_subtitle = ParagraphStyle(
        'FamSubtitle', fontName=FONT_NAME, fontSize=10,
        textColor=SECONDARY, alignment=TA_CENTER, spaceAfter=16,
    )
    style_info = ParagraphStyle(
        'FamInfo', fontName=FONT_NAME, fontSize=9,
        textColor=black, alignment=TA_LEFT, spaceAfter=3, leading=13,
    )
    style_body = ParagraphStyle(
        'FamBody', fontName=FONT_NAME, fontSize=9,
        textColor=black, alignment=TA_JUSTIFY, spaceAfter=4, leading=13,
    )
    style_h2 = ParagraphStyle(
        'FamH2', fontName=FONT_BOLD, fontSize=12, textColor=SECONDARY,
        spaceBefore=12, spaceAfter=6, leading=16,
    )
    style_h3 = ParagraphStyle(
        'FamH3', fontName=FONT_BOLD, fontSize=10, textColor=ACCENT,
        spaceBefore=8, spaceAfter=4, leading=14,
    )
    style_bullet = ParagraphStyle(
        'FamBullet', fontName=FONT_NAME, fontSize=8.5,
        textColor=HexColor("#212529"), alignment=TA_LEFT,
        spaceAfter=2, leading=12, leftIndent=12, bulletIndent=0,
    )
    style_footer = ParagraphStyle(
        'FamFooter', fontName=FONT_NAME, fontSize=7.5,
        textColor=HexColor("#6C757D"), alignment=TA_CENTER, spaceBefore=4,
    )

    now = datetime.now()
    story = []

    # Baslik
    story.append(Paragraph(_safe("Veli Bilgilendirme Ozeti"), style_title))
    story.append(Paragraph(_safe("Egitim Check-Up Psikometrik Degerlendirme"), style_subtitle))
    story.append(HRFlowable(width="100%", thickness=2, color=PRIMARY, spaceAfter=12))

    # Ogrenci bilgi tablosu
    grade_text = f"{student_grade}. Sinif" if student_grade else "-"
    info_data = [
        [Paragraph(f"<b>Ogrenci:</b> {_safe(student_name)}", style_info),
         Paragraph(f"<b>Yas:</b> {student_age}", style_info)],
        [Paragraph(f"<b>Cinsiyet:</b> {_safe(student_gender)}", style_info),
         Paragraph(f"<b>Sinif:</b> {grade_text}", style_info)],
        [Paragraph(f"<b>Tarih:</b> {now.strftime('%d.%m.%Y')}", style_info),
         Paragraph(f"<b>Testler:</b> {_safe(test_types)}", style_info)],
    ]
    info_table = Table(info_data, colWidths=[8.5 * cm, 8.5 * cm])
    info_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), HexColor("#F8F9FA")),
        ('BOX', (0, 0), (-1, -1), 0.5, BORDER),
        ('INNERGRID', (0, 0), (-1, -1), 0.25, BORDER),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
    ]))
    story.append(info_table)
    story.append(Spacer(1, 12))

    # Ozet icerigini markdown-benzeri parse et
    lines = summary_text.split('\n')
    for line in lines:
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 3))
            continue

        if stripped.startswith('### '):
            text = re.sub(r'[*#]', '', stripped[4:]).strip()
            story.append(Paragraph(f"<b>{_safe(text)}</b>", style_h2))
            continue
        if stripped.startswith('## '):
            text = re.sub(r'[*#]', '', stripped[3:]).strip()
            story.append(Paragraph(f"<b>{_safe(text)}</b>", style_h2))
            continue

        if stripped in ('---', '***', '___'):
            story.append(HRFlowable(width="100%", thickness=0.5, color=BORDER, spaceAfter=6, spaceBefore=6))
            continue

        bullet_match = re.match(r'^[-*\u2022]\s+(.+)', stripped)
        if bullet_match:
            text = _format_inline_safe(bullet_match.group(1), _safe)
            story.append(Paragraph(f"\u2022 {text}", style_bullet))
            continue

        num_match = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if num_match:
            text = _format_inline_safe(num_match.group(2), _safe)
            story.append(Paragraph(f"{num_match.group(1)}. {text}", style_bullet))
            continue

        text = _format_inline_safe(stripped, _safe)
        story.append(Paragraph(text, style_body))

    # Footer
    story.append(Spacer(1, 1 * cm))
    story.append(HRFlowable(width="100%", thickness=1, color=PRIMARY, spaceAfter=6))
    story.append(Paragraph(
        _safe("Bu belge Egitim Check-Up sistemi tarafindan olusturulmustur."),
        style_footer
    ))
    story.append(Paragraph(
        f"Olusturma tarihi: {now.strftime('%d.%m.%Y %H:%M')}",
        style_footer
    ))

    def _page_number(canvas_obj, doc_obj):
        canvas_obj.saveState()
        canvas_obj.setFont(FONT_NAME, 7)
        canvas_obj.setFillColor(HexColor("#6C757D"))
        canvas_obj.drawRightString(A4[0] - 2 * cm, 1 * cm, f"Sayfa {canvas_obj.getPageNumber()}")
        canvas_obj.restoreState()

    doc.build(story, onFirstPage=_page_number, onLaterPages=_page_number)
    buffer.seek(0)
    return buffer


def _format_inline_safe(text, safe_fn):
    """Satir ici bold/italic formatini korur."""
    s = safe_fn(text)
    s = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', s)
    s = re.sub(r'\*(.+?)\*', r'<i>\1</i>', s)
    return s


# ============================================================
# DOCX URETIMI (python-docx)
# ============================================================

def generate_family_summary_docx(student_name, student_age, student_gender,
                                 teacher_name, summary_text, test_types,
                                 student_grade=None):
    """Aile bilgilendirme ozeti icin DOCX olusturur."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    CLR_PRIMARY = RGBColor(0x1B, 0x2A, 0x4A)
    CLR_SECONDARY = RGBColor(0x2E, 0x86, 0xAB)
    CLR_GRAY = RGBColor(0x6C, 0x75, 0x7D)

    now = datetime.now()
    doc = Document()

    # Sayfa marjlari
    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2.5)

    # Baslik
    h = doc.add_heading("Veli Bilgilendirme Ozeti", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = CLR_PRIMARY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Egitim Check-Up Psikometrik Degerlendirme")
    run_sub.font.size = Pt(10)
    run_sub.font.color.rgb = CLR_SECONDARY

    doc.add_paragraph('_' * 60)

    # Ogrenci bilgileri tablosu
    grade_text = f"{student_grade}. Sinif" if student_grade else "-"
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Grid Accent 1'
    cells_data = [
        (f"Ogrenci: {student_name}", f"Yas: {student_age}"),
        (f"Cinsiyet: {student_gender}", f"Sinif: {grade_text}"),
        (f"Tarih: {now.strftime('%d.%m.%Y')}", f"Testler: {test_types}"),
    ]
    for i, (left, right) in enumerate(cells_data):
        table.cell(i, 0).text = left
        table.cell(i, 1).text = right
        for j in range(2):
            for par in table.cell(i, j).paragraphs:
                for run in par.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    # Ozet icerigini markdown-benzeri parse et
    _add_markdown_content_docx(doc, summary_text, CLR_PRIMARY, CLR_SECONDARY, CLR_GRAY)

    # Footer
    doc.add_paragraph('_' * 60)
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_foot = p_foot.add_run(
        "Bu belge Egitim Check-Up sistemi tarafindan olusturulmustur."
    )
    run_foot.font.size = Pt(7.5)
    run_foot.font.color.rgb = CLR_GRAY

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = p_date.add_run(f"Olusturma tarihi: {now.strftime('%d.%m.%Y %H:%M')}")
    run_date.font.size = Pt(7.5)
    run_date.font.color.rgb = CLR_GRAY

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _add_markdown_content_docx(doc, md_text, clr_primary, clr_secondary, clr_gray):
    """Markdown icerigini DOCX'e donusturur."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if not md_text:
        return

    for line in md_text.split('\n'):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        if stripped.startswith('### '):
            text = re.sub(r'[*#]', '', stripped[4:]).strip()
            h = doc.add_heading(text, level=3)
            for run in h.runs:
                run.font.color.rgb = clr_secondary
            continue

        if stripped.startswith('## '):
            text = re.sub(r'[*#]', '', stripped[3:]).strip()
            h = doc.add_heading(text, level=2)
            for run in h.runs:
                run.font.color.rgb = clr_primary
            continue

        if stripped in ('---', '***', '___'):
            doc.add_paragraph('_' * 40)
            continue

        bullet_match = re.match(r'^[-*\u2022]\s+(.+)', stripped)
        if bullet_match:
            _add_rich_paragraph(doc, bullet_match.group(1), bullet=True)
            continue

        num_match = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if num_match:
            _add_rich_paragraph(doc, f"{num_match.group(1)}. {num_match.group(2)}")
            continue

        _add_rich_paragraph(doc, stripped)


def _add_rich_paragraph(doc, text, bullet=False):
    """Bold (**metin**) destekli paragraf ekler."""
    from docx.shared import Pt
    import re as re_mod

    if bullet:
        p = doc.add_paragraph(style='List Bullet')
    else:
        p = doc.add_paragraph()

    parts = re_mod.split(r'(\*\*.+?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(9)
        else:
            run = p.add_run(part)
            run.font.size = Pt(9)


def generate_family_pdf_filename(student_name):
    """Aile ozeti PDF dosya adi."""
    safe_name = student_name.replace(" ", "_")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    return f"{safe_name}_Veli_Ozeti_{timestamp}.pdf"


def generate_family_docx_filename(student_name):
    """Aile ozeti DOCX dosya adi."""
    safe_name = student_name.replace(" ", "_")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    return f"{safe_name}_Veli_Ozeti_{timestamp}.docx"
